#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
地理空间文件字段批量检查工具 (GUI版本)
用于检查SHP、GDB文件及其附属文件的表字段信息

作者: ViVi141
邮箱: 747384120@qq.com
版本: 2.0 正式版
更新时间: 2025年7月14日
"""

import os
import sys
import pandas as pd
import geopandas as gpd
from pathlib import Path
import json
from datetime import datetime
import warnings
import numpy as np
import xlrd
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from shapely.geometry import Point, LineString, Polygon, MultiPolygon
from shapely.validation import make_valid
from shapely.ops import unary_union
import pyproj
from pyproj import CRS
import logging
import hashlib
import time
import psutil

# 用户友好的错误处理类
class UserFriendlyErrorHandler:
    """用户友好的错误处理类"""

    # 错误类型映射
    ERROR_TYPE_MAP = {
        'file_not_found': '文件未找到',
        'permission_denied': '权限不足',
        'encoding_error': '编码错误',
        'format_error': '格式错误',
        'memory_error': '内存不足',
        'timeout_error': '处理超时',
        'validation_error': '数据验证错误',
        'topology_error': '拓扑错误',
        'attribute_error': '属性错误',
        'geometry_error': '几何错误',
        'unknown_error': '未知错误'
    }

    # 错误解决建议
    ERROR_SOLUTIONS = {
        'file_not_found': [
            '检查文件路径是否正确',
            '确认文件是否存在于指定位置',
            '检查文件名是否包含特殊字符'
        ],
        'permission_denied': [
            '以管理员身份运行程序',
            '检查文件是否被其他程序占用',
            '确认对目标目录有读写权限'
        ],
        'encoding_error': [
            '尝试使用不同的编码格式',
            '检查文件是否损坏',
            '使用文本编辑器重新保存文件'
        ],
        'format_error': [
            '确认文件格式是否正确',
            '检查文件是否完整',
            '使用专业软件验证文件格式'
        ],
        'memory_error': [
            '关闭其他占用内存的程序',
            '分批处理大文件',
            '增加系统虚拟内存'
        ],
        'timeout_error': [
            '检查网络连接',
            '增加处理超时时间',
            '分批处理文件'
        ],
        'validation_error': [
            '检查数据是否符合标准规范',
            '修正字段类型和格式',
            '补充缺失的必填字段'
        ],
        'topology_error': [
            '修复几何要素的拓扑问题',
            '检查面要素的闭合性',
            '修正重叠和缝隙问题'
        ],
        'attribute_error': [
            '检查字段名称和类型',
            '修正字段值格式',
            '补充缺失的属性信息'
        ],
        'geometry_error': [
            '使用几何修复工具自动修复线性环未闭合问题',
            '检查几何要素的首尾点坐标是否一致',
            '使用专业GIS软件（如QGIS、ArcGIS）修复几何',
            '检查坐标精度，确保首尾点完全重合',
            '重新数字化有问题的几何要素'
        ],
        'unknown_error': [
            '重启程序后重试',
            '检查系统环境',
            '联系技术支持'
        ]
    }

    @classmethod
    def classify_error(cls, error_message):
        """分类错误类型"""
        error_lower = error_message.lower()

        if 'file not found' in error_lower or 'no such file' in error_lower:
            return 'file_not_found'
        elif 'permission denied' in error_lower or 'access denied' in error_lower:
            return 'permission_denied'
        elif 'encoding' in error_lower or 'decode' in error_lower:
            return 'encoding_error'
        elif 'format' in error_lower or 'invalid' in error_lower:
            return 'format_error'
        elif 'memory' in error_lower or 'out of memory' in error_lower:
            return 'memory_error'
        elif 'timeout' in error_lower or 'timed out' in error_lower:
            return 'timeout_error'
        elif 'validation' in error_lower or 'compliance' in error_lower:
            return 'validation_error'
        elif 'topology' in error_lower:
            return 'topology_error'
        elif 'attribute' in error_lower or 'field' in error_lower:
            return 'attribute_error'
        elif 'geometry' in error_lower or 'shape' in error_lower:
            return 'geometry_error'
        elif 'linearring' in error_lower or 'linestring' in error_lower or 'closed' in error_lower:
            return 'geometry_error'
        else:
            return 'unknown_error'

    @classmethod
    def get_user_friendly_message(cls, error_message, file_name=""):
        """获取用户友好的错误信息"""
        error_type = cls.classify_error(error_message)
        error_name = cls.ERROR_TYPE_MAP.get(error_type, '未知错误')

        # 构建用户友好的消息
        friendly_message = f"错误类型: {error_name}\n"
        if file_name:
            friendly_message += f"问题文件: {file_name}\n"
        friendly_message += f"错误详情: {error_message}\n\n"

        # 添加解决建议
        solutions = cls.ERROR_SOLUTIONS.get(error_type, [])
        if solutions:
            friendly_message += "解决建议:\n"
            for i, solution in enumerate(solutions, 1):
                friendly_message += f"{i}. {solution}\n"

        return friendly_message

    @classmethod
    def get_error_priority(cls, error_type):
        """获取错误优先级"""
        priority_map = {
            'permission_denied': 'critical',
            'memory_error': 'critical',
            'file_not_found': 'high',
            'format_error': 'high',
            'encoding_error': 'medium',
            'validation_error': 'medium',
            'topology_error': 'medium',
            'attribute_error': 'medium',
            'geometry_error': 'medium',
            'timeout_error': 'low',
            'unknown_error': 'low'
        }
        return priority_map.get(error_type, 'low')

# 导入字段编辑模块
try:
    from field_editor_dialog import FieldEditorDialog
except ImportError:
    FieldEditorDialog = None

# 导入几何编辑模块
try:
    from geometry_editor_dialog import GeometryEditorDialog
except ImportError:
    GeometryEditorDialog = None

# 配置日志系统
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('shp_checker.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# 字体配置函数
def configure_system_fonts():
    """配置系统字体"""
    try:
        from tkinter import font
        import platform

        # 根据操作系统选择合适的系统字体
        system = platform.system()
        if system == "Windows":
            # Windows系统字体
            default_font_name = "Microsoft YaHei UI"  # 微软雅黑UI
            text_font_name = "Consolas"  # 等宽字体用于代码显示
        elif system == "Darwin":  # macOS
            default_font_name = "PingFang SC"
            text_font_name = "Menlo"
        else:  # Linux
            default_font_name = "DejaVu Sans"
            text_font_name = "DejaVu Sans Mono"

        # 配置默认字体
        default_font = font.nametofont("TkDefaultFont")
        default_font.configure(family=default_font_name, size=9)

        # 配置文本字体
        text_font = font.nametofont("TkTextFont")
        text_font.configure(family=text_font_name, size=9)

        # 配置固定宽度字体
        fixed_font = font.nametofont("TkFixedFont")
        fixed_font.configure(family=text_font_name, size=9)

        logger.info(f"使用系统字体: {default_font_name} (默认), {text_font_name} (文本)")
        return True

    except Exception as e:
        logger.warning(f"字体配置失败: {e}")
        return False

# 忽略geopandas的警告
warnings.filterwarnings('ignore')

# 抑制编码转换警告
warnings.filterwarnings('ignore', category=UserWarning, module='pyogrio')
warnings.filterwarnings('ignore', category=UserWarning, module='geopandas')
warnings.filterwarnings('ignore', category=UserWarning, module='pyogrio')
warnings.filterwarnings('ignore', category=RuntimeWarning, module='pyogrio')
warnings.filterwarnings('ignore', message='.*One or several characters couldn\'t be converted correctly.*')
warnings.filterwarnings('ignore', message='.*couldn\'t be converted correctly.*')

# 默认字段标准（作为初始配置）
# 依据《中山市自然资源数据标准规范及质检规范说明》与《国土空间基础信息平台数据编目配置的质检规则》
DEFAULT_FIELD_STANDARDS = {
    # 第一组word
    "BSM": {"字段别名": "标识码", "字段类型": "Integer", "必填": True},
    "YSDM": {"字段别名": "要素代码", "字段类型": "Text", "必填": True},
    "JBNTTBBH": {"字段别名": "基本农田图斑编号", "字段类型": "Text", "必填": True},
    "TBBH": {"字段别名": "图斑编号", "字段类型": "Text", "必填": True},
    "DLBM": {"字段别名": "地类编码", "字段类型": "Text", "必填": True},
    "DLMC": {"字段别名": "地类名称", "字段类型": "Text", "必填": True},
    "QSXZ": {"字段别名": "权属性质", "字段类型": "Text", "必填": True},
    "QSDWDM": {"字段别名": "权属单位代码", "字段类型": "Text", "必填": True},
    "QSDWMC": {"字段别名": "权属单位名称", "字段类型": "Text", "必填": True},
    "ZLDWDM": {"字段别名": "坐落单位代码", "字段类型": "Text", "必填": True},
    "ZLDWMC": {"字段别名": "坐落单位名称", "字段类型": "Text", "必填": True},
    "GDLX": {"字段别名": "", "字段类型": "Text", "必填": True},
    "JBNTLX": {"字段别名": "基本农田类型", "字段类型": "Text", "必填": True},
    "ZLDJDM": {"字段别名": "质量等级代码", "字段类型": "Text", "必填": True},
    "PDJB": {"字段别名": "坡度级别", "字段类型": "Text", "必填": True},
    "KCLX": {"字段别名": "扣除类型", "字段类型": "Text", "必填": True},
    "KCDLBM": {"字段别名": "扣除地类编码", "字段类型": "Text", "必填": True},
    "TKXS": {"字段别名": "扣除地类系数", "字段类型": "Double", "必填": True},
    "XZDWMJ": {"字段别名": "线状地物面积", "字段类型": "Double", "必填": True},
    "LXDWMJ": {"字段别名": "零星地物面积", "字段类型": "Double", "必填": True},
    "TKMJ": {"字段别名": "扣除地类面积", "字段类型": "Double", "必填": True},
    "TBMJ": {"字段别名": "图斑面积", "字段类型": "Double", "必填": True},
    "JBNTMJ": {"字段别名": "基本农田面积", "字段类型": "Double", "必填": True},
    "DLBZ": {"字段别名": "地类备注", "字段类型": "Text", "字段长度": 2, "必填": True},
    # 第二组word
    "GHMC": {"字段别名": "规划名称", "字段类型": "Text", "必填": True},
    "BZ": {"字段别名": "备注", "字段类型": "Text", "必填": False},
    "CGYDMC": {"字段别名": "城规用地名称", "字段类型": "Text", "必填": True},
    "CGYDDM": {"字段别名": "城规用地代码", "字段类型": "Text", "必填": True},
    "XJYDMC": {"字段别名": "衔接用地名称", "字段类型": "Text", "必填": True},
    "CGSFQZXNR": {"字段别名": "城规是否强制性内容", "字段类型": "Text", "必填": True},
    "SSZQ": {"字段别名": "所属镇区", "字段类型": "Text", "必填": True},
    "SJSM": {"字段别名": "数据说明", "字段类型": "Text", "必填": True},
    #xlsx
    "CGLB":    {"字段别名": "城规类别", "字段类型": "Text", "必填": True, "唯一": False},
    "CSSSDY":  {"字段别名": "城市设计导引", "字段类型": "Text", "必填": False, "唯一": False},
    "DKBH":    {"字段别名": "地块编码", "字段类型": "Text", "必填": False, "唯一": False},
    "ENDTIME": {"字段别名": "ENDTIME", "字段类型": "Text", "必填": False, "唯一": False},
    "FGUID":   {"字段别名": "FGUID", "字段类型": "Text", "必填": False, "唯一": False},
    "JYSDXZ":  {"字段别名": "用地性质", "字段类型": "Text", "必填": False, "唯一": False},
    "JTCFKX":  {"字段别名": "交通出入口方位", "字段类型": "Text", "必填": False, "唯一": False},
    "JZMDX":   {"字段别名": "建筑密度下限", "字段类型": "Double", "必填": True, "唯一": False},
    "JZMDZ":   {"字段别名": "建筑密度上限", "字段类型": "Double", "必填": True, "唯一": False},
    "JZXG":    {"字段别名": "建筑限高", "字段类型": "Double", "必填": True, "唯一": False},
    "KGLSX":   {"字段别名": "控规绿地率下限", "字段类型": "Double", "必填": True, "唯一": False},
    "LDLSX":   {"字段别名": "绿地率上限", "字段类型": "Double", "必填": True, "唯一": False},
    "LDLLX":   {"字段别名": "绿地率下限", "字段类型": "Double", "必填": True, "唯一": False},
    "NJLJXZL": {"字段别名": "年均流总量控制率", "字段类型": "Double", "必填": True, "唯一": False},
    "PFDATE":  {"字段别名": "批准日期", "字段类型": "Text", "必填": True, "唯一": False},
    "PFNAME":  {"字段别名": "批准文号", "字段类型": "Text", "必填": False, "唯一": False},
    "PTGSS":   {"字段别名": "配套公共实施项目与规模", "字段类型": "Text", "必填": True, "唯一": False},
    "PTSZS":   {"字段别名": "配套市政设施", "字段类型": "Text", "必填": True, "唯一": False},
    "RJLSX":   {"字段别名": "容积率上限", "字段类型": "Double", "必填": True, "唯一": False},
    "RJLXX":   {"字段别名": "容积率下限", "字段类型": "Double", "必填": True, "唯一": False},
    "SFQZXNR": {"字段别名": "是否强制性内容", "字段类型": "Text", "必填": True, "唯一": False},
    "STARTTIME": {"字段别名": "STARTTIME", "字段类型": "日期", "必填": True, "唯一": False},
    "SUOSZQ":  {"字段别名": "所属镇区", "字段类型": "Text", "必填": False, "唯一": False},
    "TCW":     {"字段别名": "停车位", "字段类型": "Double", "必填": True, "唯一": False},
    "TSZPLTJZ": {"字段别名": "透水砖铺装率下限", "字段类型": "Double", "必填": True, "唯一": False},
    "TSZPLZDJZ": {"字段别名": "透水砖铺装率上限", "字段类型": "Double", "必填": True, "唯一": False},
    "XCSLDLTJZ": {"字段别名": "下沉式绿地率上限", "字段类型": "Double", "必填": True, "唯一": False},
    "XCSLDLZDJZ": {"字段别名": "下沉式绿地率下限", "字段类型": "Double", "必填": True, "唯一": False},
    "XJYDLB":  {"字段别名": "衔接用地类别", "字段类型": "Text", "必填": False, "唯一": False},
    "YDLX":    {"字段别名": "用地类型", "字段类型": "Text", "必填": False, "唯一": False},
    "YDMJ":    {"字段别名": "用地面积", "字段类型": "Double", "必填": False, "唯一": False},
    "YDXZDM":  {"字段别名": "用地性质代码", "字段类型": "Text", "必填": False, "唯一": False},
    "YDXZMC":  {"字段别名": "用地性质名称", "字段类型": "Text", "必填": False, "唯一": False},
    "ZMJ":     {"字段别名": "总建筑面积", "字段类型": "Double", "必填": False, "唯一": False},
    "ZQCODE":  {"字段别名": "镇区编码", "字段类型": "Text", "必填": False, "唯一": False},
}

# 当前字段标准（可动态修改）
FIELD_STANDARDS = DEFAULT_FIELD_STANDARDS.copy()

# 字段类型映射
FIELD_TYPE_MAP = {
    '文本': 'object', 'Text': 'object',
    '双精度': 'float', 'Double': 'float',
    '整数': 'int', 'Integer': 'int',
    '日期': 'datetime', 'Date': 'datetime', 'Datetime': 'datetime',
}

# 错误等级定义
ERROR_LEVELS = {
    'IGNORABLE': '可忽略',
    'CRITICAL': '不可忽略'
}

# 错误类型定义
ERROR_TYPES = {
    'ENCODING_ERROR': '编码错误',
    'GEOMETRY_ERROR': '几何错误',
    'FIELD_TYPE_ERROR': '字段类型错误',
    'REQUIRED_FIELD_ERROR': '必填字段错误',
    'TOPOLOGY': '拓扑错误',
    'ATTRIBUTE': '属性错误',
    'BASIC': '基础错误',
    'DATA_INTEGRITY': '数据完整性错误',
    'LOGICAL_CONSISTENCY': '逻辑一致性错误',
    'SPATIAL_REFERENCE': '空间参考错误',
    'FIELD_VALUE_CONSISTENCY': '字段值一致性错误',
    'OTHER_ERROR': '其他错误'
}



# 编号唯一性字段
UNIQUE_FIELDS = ['BSM', 'TBBH', 'JBNTTBBH', 'DKBH', 'FGUID']

def calculate_file_hash(file_path: Path, algorithm: str = 'sha256') -> str:
    """计算文件的哈希值

    Args:
        file_path: 文件路径
        algorithm: 哈希算法，默认为sha256

    Returns:
        文件的哈希值字符串
    """
    try:
        hash_obj = hashlib.new(algorithm)
        with open(file_path, 'rb') as f:
            # 分块读取大文件，避免内存溢出
            for chunk in iter(lambda: f.read(4096), b""):
                hash_obj.update(chunk)
        return hash_obj.hexdigest()
    except Exception as e:
        logger.error(f"计算文件哈希值失败 {file_path}: {e}")
        return "计算失败"

class FieldConfigManager:
    """字段配置管理器

    默认字段标准依据《中山市自然资源数据标准规范及质检规范说明》
    与《国土空间基础信息平台数据编目配置的质检规则》制定
    """

    def __init__(self, config_file="field_config.json"):
        self.config_file = config_file
        self.field_standards = DEFAULT_FIELD_STANDARDS.copy()
        self.load_config()

    def load_config(self):
        """加载配置文件"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    loaded_data = json.load(f)

                # 验证数据格式
                if isinstance(loaded_data, dict):
                    # 验证每个字段的配置格式
                    valid_data = {}
                    for field_name, field_config in loaded_data.items():
                        if isinstance(field_config, dict):
                            # 确保必要的字段存在
                            valid_config = {
                                "字段别名": field_config.get("字段别名", field_name),
                                "字段类型": field_config.get("字段类型", "Text"),
                                "必填": field_config.get("必填", False),
                                "唯一": field_config.get("唯一", False)
                            }
                            # 添加可选字段
                            if "字段长度" in field_config:
                                valid_config["字段长度"] = field_config["字段长度"]

                            valid_data[field_name] = valid_config

                    self.field_standards = valid_data
                    logger.info(f"已加载字段配置文件: {self.config_file}，包含 {len(valid_data)} 个字段")
                else:
                    raise ValueError("配置文件格式错误")
            else:
                logger.info("配置文件不存在，使用默认配置")
                self.field_standards = DEFAULT_FIELD_STANDARDS.copy()
        except Exception as e:
            logger.error(f"加载配置文件失败: {e}")
            logger.info("使用默认配置")
            self.field_standards = DEFAULT_FIELD_STANDARDS.copy()

    def save_config(self):
        """保存配置文件"""
        try:
            # 创建备份
            backup_file = self.config_file + ".backup"
            if os.path.exists(self.config_file):
                import shutil
                shutil.copy2(self.config_file, backup_file)
                logger.info(f"已创建备份文件: {backup_file}")

            # 保存配置
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.field_standards, f, ensure_ascii=False, indent=2)

            logger.info(f"已保存字段配置文件: {self.config_file}")
            return True
        except Exception as e:
            logger.error(f"保存配置文件失败: {e}")
            return False

    def get_field_standards(self):
        """获取当前字段标准"""
        return self.field_standards

    def update_field_standards(self, new_standards):
        """更新字段标准"""
        self.field_standards = new_standards.copy()
        # 同时更新全局变量
        global FIELD_STANDARDS
        FIELD_STANDARDS = self.field_standards.copy()

    def add_field(self, field_name, field_config):
        """添加字段"""
        self.field_standards[field_name] = field_config
        global FIELD_STANDARDS
        FIELD_STANDARDS = self.field_standards.copy()

    def remove_field(self, field_name):
        """删除字段"""
        if field_name in self.field_standards:
            del self.field_standards[field_name]
            global FIELD_STANDARDS
            FIELD_STANDARDS = self.field_standards.copy()

    def reset_to_default(self):
        """重置为默认配置"""
        self.field_standards = DEFAULT_FIELD_STANDARDS.copy()
        global FIELD_STANDARDS
        FIELD_STANDARDS = self.field_standards.copy()

def summarize_required_field_errors(required_field_issues):
    """汇总必填字段错误，优化显示同一行的多个错误"""
    if not required_field_issues:
        return []

    # 按行索引分组
    row_errors = {}
    critical_errors = 0
    ignorable_errors = 0

    for issue in required_field_issues:
        row_idx = issue.get('row_index', 0)
        if row_idx not in row_errors:
            row_errors[row_idx] = {'fields': [], 'critical': 0, 'ignorable': 0}

        # 统计错误等级
        for field_error in issue.get('field_errors', []):
            field_name = field_error['field_name']
            error_level = field_error['error_level']
            row_errors[row_idx]['fields'].append(field_name)

            if error_level == ERROR_LEVELS['CRITICAL']:
                row_errors[row_idx]['critical'] += 1
                critical_errors += 1
            else:
                row_errors[row_idx]['ignorable'] += 1
                ignorable_errors += 1

    # 生成汇总报告
    summary_issues = []

    # 统计信息
    total_rows_with_errors = len(row_errors)
    all_missing_fields = set()
    for row_data in row_errors.values():
        all_missing_fields.update(row_data['fields'])

    # 添加总体统计
    summary_issues.append({
        'type': '必填字段错误汇总',
        'error': f'共有{total_rows_with_errors}行存在必填字段缺失，涉及字段: {", ".join(sorted(all_missing_fields))}',
        'detail': f'不可忽略错误: {critical_errors}个，可忽略错误: {ignorable_errors}个',
        'total_rows': total_rows_with_errors,
        'total_fields': len(all_missing_fields),
        'critical_errors': critical_errors,
        'ignorable_errors': ignorable_errors
    })

    # 按错误等级分组
    critical_rows = [row_idx for row_idx, data in row_errors.items() if data['critical'] > 0]
    ignorable_rows = [row_idx for row_idx, data in row_errors.items() if data['ignorable'] > 0]

    if critical_rows:
        summary_issues.append({
            'type': '不可忽略错误',
            'error': f'有{len(critical_rows)}行存在不可忽略的必填字段缺失',
            'detail': f'涉及行: {", ".join([str(r+1) for r in critical_rows])}',
            'error_level': ERROR_LEVELS['CRITICAL']
        })

    if ignorable_rows:
        summary_issues.append({
            'type': '可忽略错误',
            'error': f'有{len(ignorable_rows)}行存在可忽略的必填字段缺失',
            'detail': f'涉及行: {", ".join([str(r+1) for r in ignorable_rows])}',
            'error_level': ERROR_LEVELS['IGNORABLE']
        })

    # 添加最严重的错误行（不可忽略错误最多的行）
    if critical_rows:
        max_critical_row = max(critical_rows, key=lambda r: row_errors[r]['critical'])
        critical_fields = [f for f in row_errors[max_critical_row]['fields']
                         if any(fe['field_name'] == f and fe['error_level'] == ERROR_LEVELS['CRITICAL']
                               for fe in next(issue for issue in required_field_issues
                                            if issue.get('row_index') == max_critical_row).get('field_errors', []))]
        summary_issues.append({
            'type': '最严重错误行',
            'error': f'第{max_critical_row+1}行存在{row_errors[max_critical_row]["critical"]}个不可忽略错误',
            'detail': f'不可忽略字段: {", ".join(critical_fields)}',
            'error_level': ERROR_LEVELS['CRITICAL']
        })

    return summary_issues

def get_field_error_level(field_name, file_name):
    """根据字段名和文件名确定错误等级"""
    # 转换为大写以便比较
    file_name_upper = file_name.upper()

    # 定义特殊规则（默认配置，可被用户配置覆盖）
    critical_fields_for_special_files = {
        'GHMC': ['YDFW', 'GHJX'],  # 规划名称字段在YDFW或GHJX文件中为不可忽略
        'PFDATE': ['YDFW', 'GHJX']  # 批准日期字段在YDFW或GHJX文件中为不可忽略
    }

    # 检查是否为特殊字段
    if field_name in critical_fields_for_special_files:
        required_patterns = critical_fields_for_special_files[field_name]
        for pattern in required_patterns:
            if pattern in file_name_upper:
                return ERROR_LEVELS['CRITICAL']  # 不可忽略

    # 默认返回可忽略
    return ERROR_LEVELS['IGNORABLE']

def check_required_fields_detailed(gdf, field_standards, file_name=None):
    """详细检查必填字段，返回具体的空值行信息"""
    issues = []

    # 获取所有必填字段
    required_fields = []
    for field_name, standard in field_standards.items():
        if standard.get('必填') or str(standard.get('约束条件', '')).strip().upper() == 'O':
            if field_name in gdf.columns:
                required_fields.append(field_name)

    # 添加调试信息
    total_required_in_standard = len([f for f, s in field_standards.items() if s.get('必填')])
    missing_required_fields = [f for f, s in field_standards.items() if s.get('必填') and f not in gdf.columns]

    logger.info(f"检查必填字段: 标准中定义了{total_required_in_standard}个必填字段")
    logger.info(f"数据文件中存在{len(required_fields)}个必填字段: {required_fields}")
    logger.info(f"数据文件所有字段: {list(gdf.columns)}")

    if missing_required_fields:
        logger.info(f"标准中定义但数据文件中不存在的必填字段: {missing_required_fields}")

    if not required_fields:
        logger.warning("未找到任何必填字段，请检查字段标准配置")
        return issues

    # 检查每一行的必填字段
    total_rows = len(gdf)
    rows_with_errors = 0

    for row_idx in range(total_rows):
        missing_fields = []
        for field_name in required_fields:
            # 更严格的空值检查
            value = gdf.iloc[row_idx][field_name]
            if pd.isna(value) or value == '' or str(value).strip() == '':
                missing_fields.append(field_name)

        if missing_fields:
            rows_with_errors += 1

            # 为每个缺失字段确定错误等级
            field_errors = []
            for field_name in missing_fields:
                error_level = get_field_error_level(field_name, file_name or '')
                field_errors.append({
                    'field_name': field_name,
                    'error_level': error_level,
                    'error_type': 'REQUIRED_FIELD_ERROR'
                })

            issues.append({
                'row_index': row_idx,
                'missing_fields': missing_fields,
                'field_errors': field_errors,
                'error': f'第{row_idx+1}行缺少必填字段: {", ".join(missing_fields)}',
                'type': '必填字段错误'
            })

    # 添加统计信息
    if rows_with_errors > 0:
        logger.info(f"必填字段检查完成: 总共{total_rows}行，其中{rows_with_errors}行存在必填字段缺失")

        # 统计每个字段的缺失情况
        field_missing_stats = {}
        critical_errors = 0
        ignorable_errors = 0

        for issue in issues:
            for field_error in issue.get('field_errors', []):
                field_name = field_error['field_name']
                error_level = field_error['error_level']

                if field_name not in field_missing_stats:
                    field_missing_stats[field_name] = {'critical': 0, 'ignorable': 0}

                if error_level == ERROR_LEVELS['CRITICAL']:
                    field_missing_stats[field_name]['critical'] += 1
                    critical_errors += 1
                else:
                    field_missing_stats[field_name]['ignorable'] += 1
                    ignorable_errors += 1

        for field, stats in field_missing_stats.items():
            if stats['critical'] > 0:
                logger.info(f"字段 {field} 不可忽略错误 {stats['critical']} 次")
            if stats['ignorable'] > 0:
                logger.info(f"字段 {field} 可忽略错误 {stats['ignorable']} 次")

        logger.info(f"总计: 不可忽略错误 {critical_errors} 个，可忽略错误 {ignorable_errors} 个")

    # 如果错误数量较多，进行汇总
    if len(issues) > 10:
        return summarize_required_field_errors(issues)
    else:
        return issues

def check_field_compliance(field_name, series, standard):
    """检查单字段合规性，返回问题列表"""
    issues = []

    # 类型检查
    std_type = FIELD_TYPE_MAP.get(str(standard.get('字段类型', '')).strip(), None)
    if std_type:
        if std_type == 'object' and not (series.dtype == 'object' or str(series.dtype).startswith('str')):
            issues.append(f"类型不符，应为文本，实际为{series.dtype}")
        elif std_type == 'float' and not (str(series.dtype).startswith('float') or str(series.dtype).startswith('double')):
            issues.append(f"类型不符，应为双精度，实际为{series.dtype}")
        elif std_type == 'int' and not (str(series.dtype).startswith('int')):
            issues.append(f"类型不符，应为整数，实际为{series.dtype}")
        elif std_type == 'datetime' and not (str(series.dtype).startswith('datetime')):
            issues.append(f"类型不符，应为日期，实际为{series.dtype}")

    # 必填检查（简化版本，详细检查在check_required_fields_detailed中）
    if standard.get('必填') or str(standard.get('约束条件', '')).strip().upper() == 'O':
        null_count = series.isnull().sum()
        if null_count > 0:
            issues.append(f"必填字段存在空值，共{null_count}个")

    # 字段长度检查
    if '字段长度' in standard and standard['字段长度']:
        max_length = standard['字段长度']
        if series.dtype == 'object':
            # 检查文本字段长度
            max_str_length = series.astype(str).str.len().max()
            if max_str_length > max_length:
                issues.append(f"字段长度超限，最大长度{max_str_length}，限制为{max_length}")

    # 数值范围检查（针对特定字段）
    if std_type == 'float' or std_type == 'int':
        if field_name in ['JZMDX', 'JZMDZ', 'JZXG', 'KGLSX', 'LDLSX', 'LDLLX', 'NJLJXZL', 'RJLSX', 'RJLXX', 'TCW', 'TSZPLTJZ', 'TSZPLZDJZ', 'XCSLDLTJZ', 'XCSLDLZDJZ']:
            # 检查数值是否在合理范围内
            non_null_values = series.dropna()
            if len(non_null_values) > 0:
                min_val = non_null_values.min()
                max_val = non_null_values.max()

                # 根据字段类型设置合理范围
                if field_name in ['JZMDX', 'JZMDZ']:  # 建筑密度
                    if min_val < 0 or max_val > 100:
                        issues.append(f"建筑密度值超出合理范围[0-100]，实际范围[{min_val}-{max_val}]")
                elif field_name in ['JZXG']:  # 建筑限高
                    if min_val < 0 or max_val > 1000:
                        issues.append(f"建筑限高值超出合理范围[0-1000]，实际范围[{min_val}-{max_val}]")
                elif field_name in ['KGLSX', 'LDLSX', 'LDLLX']:  # 绿地率
                    if min_val < 0 or max_val > 100:
                        issues.append(f"绿地率值超出合理范围[0-100]，实际范围[{min_val}-{max_val}]")
                elif field_name in ['RJLSX', 'RJLXX']:  # 容积率
                    if min_val < 0 or max_val > 50:
                        issues.append(f"容积率值超出合理范围[0-50]，实际范围[{min_val}-{max_val}]")
                elif field_name in ['TCW']:  # 停车位
                    if min_val < 0:
                        issues.append(f"停车位数量不能为负数，最小值{min_val}")

    # 编码格式检查
    if std_type == 'object' and field_name in ['YSDM', 'DLBM', 'QSDWDM', 'ZLDWDM', 'ZLDJDM', 'PDJB', 'KCLX', 'KCDLBM', 'CGYDDM', 'YDXZDM', 'ZQCODE']:
        # 检查编码字段是否包含非法字符
        non_null_values = series.dropna().astype(str)
        if len(non_null_values) > 0:
            # 检查是否包含中文字符（编码字段通常不应包含中文）
            chinese_chars = non_null_values.str.contains(r'[\u4e00-\u9fff]', na=False)
            if chinese_chars.any():
                issues.append(f"编码字段包含中文字符，可能影响数据规范性")

            # 检查是否包含特殊字符
            special_chars = non_null_values.str.contains(r'[^\w\-\.]', na=False)
            if special_chars.any():
                issues.append(f"编码字段包含特殊字符，可能影响数据规范性")

    return issues

def check_topology_gaps(geometries, tolerance=0.001):
    """检查面缝隙 - 使用改进的算法"""
    try:
        # 尝试使用改进的算法
        from improved_topology_utils import ImprovedTopologyChecker
        checker = ImprovedTopologyChecker(tolerance)
        gaps = checker.check_topology_gaps_optimized(geometries, tolerance)

        # 转换为原始格式以保持兼容性
        legacy_gaps = []
        for gap in gaps:
            legacy_gaps.append({
                'feature1': gap['feature1'],
                'feature2': gap['feature2'],
                'distance': gap['distance'],
                'type': '面缝隙'
            })
        return legacy_gaps

    except ImportError:
        # 如果改进模块不可用，使用原始算法
        logger.warning("改进的拓扑检查模块不可用，使用原始算法")
        return check_topology_gaps_original(geometries, tolerance)
    except Exception as e:
        logger.error(f"改进算法失败，回退到原始算法: {e}")
        return check_topology_gaps_original(geometries, tolerance)

def check_topology_gaps_original(geometries, tolerance=0.001):
    """原始面缝隙检测算法（备用）"""
    gaps = []
    for i, geom1 in enumerate(geometries):
        if geom1 is None or geom1.is_empty:
            continue
        for j, geom2 in enumerate(geometries[i+1:], i+1):
            if geom2 is None or geom2.is_empty:
                continue
            try:
                # 检查两个几何体之间的距离
                distance = geom1.distance(geom2)
                if 0 < distance < tolerance:
                    gaps.append({
                        'feature1': i,
                        'feature2': j,
                        'distance': distance,
                        'type': '面缝隙'
                    })
            except Exception as e:
                continue
    return gaps

def check_topology_overlaps(geometries, tolerance=0.001):
    """检查面重叠"""
    overlaps = []
    for i, geom1 in enumerate(geometries):
        if geom1 is None or geom1.is_empty:
            continue
        for j, geom2 in enumerate(geometries[i+1:], i+1):
            if geom2 is None or geom2.is_empty:
                continue
            try:
                # 检查两个几何体是否重叠
                if geom1.overlaps(geom2):
                    intersection = geom1.intersection(geom2)
                    if hasattr(intersection, 'area') and intersection.area > tolerance:
                        overlaps.append({
                            'feature1': i,
                            'feature2': j,
                            'overlap_area': intersection.area,
                            'type': '面重叠'
                        })
            except Exception as e:
                continue
    return overlaps

def check_geometry_validity(geometries, auto_fix=False, tolerance=0.001):
    """检查几何有效性，可选择自动修复

    Args:
        geometries: 几何对象列表
        auto_fix: 是否自动修复可修复的几何错误
        tolerance: 修复容差

    Returns:
        invalid_geometries: 无效几何列表
        fixed_geometries: 已修复的几何索引列表（如果auto_fix=True）
    """
    invalid_geometries = []
    fixed_geometries = []

    for i, geom in enumerate(geometries):
        if geom is None:
            invalid_geometries.append({
                'feature': i,
                'error': '几何为空',
                'type': '几何检查',
                'severity': 'critical'
            })
        elif geom.is_empty:
            invalid_geometries.append({
                'feature': i,
                'error': '几何为空几何',
                'type': '几何检查',
                'severity': 'critical'
            })
        elif not geom.is_valid:
            # 尝试修复几何
            try:
                fixed_geom = make_valid(geom)
                if fixed_geom.is_valid:
                    if auto_fix:
                        # 自动修复模式：直接替换原几何
                        geometries[i] = fixed_geom
                        fixed_geometries.append(i)
                        logger.info(f"已自动修复几何 {i}: {geom.is_valid_reason if hasattr(geom, 'is_valid_reason') else '线性环未闭合'}")
                    else:
                        # 检测模式：只记录问题
                        invalid_geometries.append({
                            'feature': i,
                            'error': f'几何无效但可修复: {geom.is_valid_reason if hasattr(geom, "is_valid_reason") else "线性环未闭合"}',
                            'type': '几何检查',
                            'severity': 'fixable',
                            'original_error': geom.is_valid_reason if hasattr(geom, "is_valid_reason") else "线性环未闭合",
                            'fix_suggestion': '使用几何修复工具自动修复'
                        })
                else:
                    invalid_geometries.append({
                        'feature': i,
                        'error': f'几何无效且无法修复: {geom.is_valid_reason if hasattr(geom, "is_valid_reason") else "未知原因"}',
                        'type': '几何检查',
                        'severity': 'critical'
                    })
            except Exception as e:
                invalid_geometries.append({
                    'feature': i,
                    'error': f'几何无效且修复失败: {geom.is_valid_reason if hasattr(geom, "is_valid_reason") else "未知原因"} - {str(e)}',
                    'type': '几何检查',
                    'severity': 'critical'
                })

    return invalid_geometries, fixed_geometries

def auto_fix_geometry_file(file_path, tolerance=0.001):
    """自动修复单个文件的几何错误

    Args:
        file_path: 文件路径
        tolerance: 修复容差

    Returns:
        dict: 修复结果统计
    """
    try:
        logger.info(f"开始自动修复几何文件: {file_path}")

        # 读取文件
        gdf = gpd.read_file(file_path)
        if gdf.empty:
            return {'success': False, 'error': '文件为空'}

        # 统计修复前的几何问题
        total_geometries = len(gdf)
        invalid_before = sum(1 for geom in gdf.geometry if geom is not None and not geom.is_valid)

        if invalid_before == 0:
            return {'success': True, 'message': '没有发现几何错误', 'fixed_count': 0}

        # 修复几何错误
        fixed_count = 0
        error_count = 0
        fixed_indices = []

        for idx, row in gdf.iterrows():
            try:
                geom = row.geometry
                if geom is not None and not geom.is_valid:
                    # 尝试修复
                    fixed_geom = make_valid(geom)
                    if fixed_geom.is_valid:
                        gdf.at[idx, 'geometry'] = fixed_geom
                        fixed_count += 1
                        fixed_indices.append(idx)
                        logger.info(f"已修复几何 {idx}")
                    else:
                        error_count += 1
                        logger.warning(f"几何 {idx} 无法修复")
            except Exception as e:
                error_count += 1
                logger.error(f"修复几何 {idx} 时出错: {e}")

        # 保存修复后的文件
        if fixed_count > 0:
            # 创建备份文件
            backup_path = str(file_path) + '.backup'
            gdf.to_file(backup_path)
            logger.info(f"已创建备份文件: {backup_path}")

            # 保存修复后的文件
            gdf.to_file(file_path)
            logger.info(f"已保存修复后的文件: {file_path}")

        result = {
            'success': True,
            'total_geometries': total_geometries,
            'invalid_before': invalid_before,
            'fixed_count': fixed_count,
            'error_count': error_count,
            'fixed_indices': fixed_indices,
            'backup_path': backup_path if fixed_count > 0 else None
        }

        logger.info(f"几何修复完成: 修复 {fixed_count} 个，失败 {error_count} 个")
        return result

    except Exception as e:
        error_msg = f"自动修复几何文件失败: {str(e)}"
        logger.error(error_msg)
        return {'success': False, 'error': error_msg}

def auto_fix_geometry_batch(file_paths, tolerance=0.001, progress_callback=None):
    """批量自动修复多个文件的几何错误

    Args:
        file_paths: 文件路径列表
        tolerance: 修复容差
        progress_callback: 进度回调函数

    Returns:
        dict: 批量修复结果统计
    """
    total_files = len(file_paths)
    successful_files = 0
    failed_files = 0
    total_fixed = 0
    total_errors = 0
    results = []

    logger.info(f"开始批量修复 {total_files} 个文件的几何错误")

    for i, file_path in enumerate(file_paths):
        try:
            if progress_callback:
                progress_callback(i + 1, total_files, f"正在修复: {Path(file_path).name}")

            # 修复单个文件
            result = auto_fix_geometry_file(file_path, tolerance)
            results.append({
                'file_path': file_path,
                'result': result
            })

            if result['success']:
                successful_files += 1
                if 'fixed_count' in result:
                    total_fixed += result['fixed_count']
                if 'error_count' in result:
                    total_errors += result['error_count']
            else:
                failed_files += 1

        except Exception as e:
            failed_files += 1
            error_msg = f"处理文件 {file_path} 时出错: {str(e)}"
            logger.error(error_msg)
            results.append({
                'file_path': file_path,
                'result': {'success': False, 'error': error_msg}
            })

    summary = {
        'total_files': total_files,
        'successful_files': successful_files,
        'failed_files': failed_files,
        'total_fixed': total_fixed,
        'total_errors': total_errors,
        'results': results
    }

    logger.info(f"批量修复完成: 成功 {successful_files} 个文件，失败 {failed_files} 个文件，共修复 {total_fixed} 个几何错误")
    return summary

def check_coordinate_system(gdf):
    """检查数学基础（坐标系统）"""
    issues = []

    # 检查CRS是否存在
    if gdf.crs is None:
        issues.append({
            'type': '数学基础',
            'error': '缺少坐标参考系统(CRS)'
        })
    else:
        # 检查CRS是否有效
        try:
            crs_info = gdf.crs.to_string()
            if not crs_info:
                issues.append({
                    'type': '数学基础',
                    'error': '坐标参考系统信息无效'
                })
        except Exception as e:
            issues.append({
                'type': '数学基础',
                'error': f'坐标参考系统错误: {str(e)}'
            })

    return issues

def check_attribute_structure_consistency(dataframes):
    """检查属性结构一致性 - 只检查标准字段的类型一致性"""
    issues = []
    if not dataframes:
        return issues

    # 收集所有数据框中标准字段的类型信息
    all_field_types = {}

    for i, df in enumerate(dataframes):
        for col in df.columns:
            # 只检查标准字段定义中的字段
            if col in FIELD_STANDARDS and col != 'geometry':
                if col not in all_field_types:
                    all_field_types[col] = {}
                all_field_types[col][i] = df[col].dtype

    # 检查同一字段在不同文件中的类型一致性
    for field_name, file_types in all_field_types.items():
        if len(file_types) > 1:  # 只检查在多个文件中都存在的字段
            type_values = list(file_types.values())
            first_type = type_values[0]

            # 检查是否有类型不一致的情况
            for file_index, dtype in file_types.items():
                if dtype != first_type:
                    issues.append({
                        'file_index': file_index,
                        'type': '属性结构一致性',
                        'error': f'字段 {field_name} 类型不一致: 应为{first_type}, 实际为{dtype}'
                    })

    return issues

def check_numeric_ranges(gdf):
    """检查数值范围符合性 - 已取消数值范围定义"""
    # 数值范围检查已取消
    return []

def check_unique_identifiers(dataframes):
    """检查编号唯一性（单文件内）- 只检测表中实际存在的字段"""
    issues = []

    # 对每个文件单独检查唯一性
    for file_index, df in enumerate(dataframes):
        # 获取该文件中实际存在的唯一性字段
        existing_unique_fields = []
        for field_name in UNIQUE_FIELDS:
            if field_name in df.columns:
                existing_unique_fields.append(field_name)

        # 检查该文件中的唯一性
        for field_name in existing_unique_fields:
            values = df[field_name].dropna().astype(str).tolist()

            if values:
                # 检查重复值
                value_counts = pd.Series(values).value_counts()

                for value, count in value_counts.items():
                    if count > 1:
                        issues.append({
                            'field': field_name,
                            'type': '编号唯一性',
                            'error': f'字段 {field_name} 的值 "{value}" 在文件内重复出现 {count} 次',
                            'duplicate_value': value,
                            'duplicate_count': count,
                            'file_index': file_index
                        })

    return issues

def check_data_integrity(gdf):
    """检查数据完整性"""
    issues = []

    # 检查几何数据完整性
    if not gdf.empty:
        # 检查是否有空几何
        null_geometries = gdf.geometry.isnull().sum()
        if null_geometries > 0:
            issues.append({
                'type': '数据完整性',
                'error': f'存在{null_geometries}个空几何对象'
            })

        # 检查是否有空几何体
        empty_geometries = gdf.geometry.apply(lambda x: x.is_empty if x is not None else False).sum()
        if empty_geometries > 0:
            issues.append({
                'type': '数据完整性',
                'error': f'存在{empty_geometries}个空几何体'
            })

    return issues

def check_logical_consistency(gdf):
    """检查逻辑一致性"""
    issues = []

    if not gdf.empty:
        # 检查面积字段逻辑一致性
        area_fields = ['TBMJ', 'JBNTMJ', 'XZDWMJ', 'LXDWMJ', 'TKMJ', 'YDMJ', 'ZMJ']
        existing_area_fields = [field for field in area_fields if field in gdf.columns]

        if len(existing_area_fields) >= 2:
            # 检查面积字段之间的逻辑关系
            for i, field1 in enumerate(existing_area_fields):
                for field2 in existing_area_fields[i+1:]:
                    if field1 in gdf.columns and field2 in gdf.columns:
                        # 检查是否有负面积
                        neg_area1 = (gdf[field1] < 0).sum()
                        neg_area2 = (gdf[field2] < 0).sum()

                        if neg_area1 > 0:
                            issues.append({
                                'type': '逻辑一致性',
                                'error': f'字段{field1}存在{neg_area1}个负面积值'
                            })

                        if neg_area2 > 0:
                            issues.append({
                                'type': '逻辑一致性',
                                'error': f'字段{field2}存在{neg_area2}个负面积值'
                            })

        # 检查编码字段格式一致性
        code_fields = ['YSDM', 'DLBM', 'QSDWDM', 'ZLDWDM']
        for field in code_fields:
            if field in gdf.columns:
                # 检查编码长度是否一致
                non_null_codes = gdf[field].dropna().astype(str)
                if len(non_null_codes) > 0:
                    code_lengths = non_null_codes.str.len()
                    if code_lengths.nunique() > 1:
                        issues.append({
                            'type': '逻辑一致性',
                            'error': f'字段{field}编码长度不一致，长度范围[{code_lengths.min()}-{code_lengths.max()}]'
                        })

    return issues

def check_spatial_reference_consistency(gdf):
    """检查空间参考一致性"""
    issues = []

    if not gdf.empty and gdf.crs is not None:
        try:
            crs_string = gdf.crs.to_string()

            # 检查是否为常用坐标系
            common_crs = [
                'EPSG:4326',  # WGS84
                'EPSG:3857',  # Web Mercator
                'EPSG:4490',  # CGCS2000
                'EPSG:4547',  # CGCS2000 / 3-degree Gauss-Kruger zone 39
                'EPSG:4548',  # CGCS2000 / 3-degree Gauss-Kruger zone 40
                'EPSG:4549',  # CGCS2000 / 3-degree Gauss-Kruger zone 41
                'EPSG:4550',  # CGCS2000 / 3-degree Gauss-Kruger zone 42
                'EPSG:4551',  # CGCS2000 / 3-degree Gauss-Kruger zone 43
                'EPSG:4552',  # CGCS2000 / 3-degree Gauss-Kruger zone 44
                'EPSG:4553',  # CGCS2000 / 3-degree Gauss-Kruger zone 45
            ]

            if not any(crs in crs_string for crs in common_crs):
                issues.append({
                    'type': '空间参考一致性',
                    'error': f'使用了非标准坐标系: {crs_string}'
                })

            # 检查坐标范围是否合理
            bounds = gdf.total_bounds
            if bounds is not None:
                min_x, min_y, max_x, max_y = bounds

                # 检查是否在中国范围内（大致范围）
                if not (73 <= min_x <= 135 and 18 <= min_y <= 54):
                    issues.append({
                        'type': '空间参考一致性',
                        'error': f'坐标范围超出中国范围，当前范围: X[{min_x:.6f}-{max_x:.6f}], Y[{min_y:.6f}-{max_y:.6f}]'
                    })

                # 检查坐标精度
                if abs(max_x - min_x) < 0.000001 or abs(max_y - min_y) < 0.000001:
                    issues.append({
                        'type': '空间参考一致性',
                        'error': '坐标范围过小，可能存在坐标精度问题'
                    })

        except Exception as e:
            issues.append({
                'type': '空间参考一致性',
                'error': f'坐标系统检查失败: {str(e)}'
            })

    return issues

def check_field_value_consistency(gdf):
    """检查字段值一致性"""
    issues = []

    if not gdf.empty:
        # 检查地类编码与地类名称的一致性
        if 'DLBM' in gdf.columns and 'DLMC' in gdf.columns:
            # 检查是否有地类编码但无地类名称的情况
            has_code_no_name = ((gdf['DLBM'].notna()) & (gdf['DLMC'].isna())).sum()
            if has_code_no_name > 0:
                issues.append({
                    'type': '字段值一致性',
                    'error': f'存在{has_code_no_name}条记录有地类编码但无地类名称'
                })

            # 检查是否有地类名称但无地类编码的情况
            has_name_no_code = ((gdf['DLMC'].notna()) & (gdf['DLBM'].isna())).sum()
            if has_name_no_code > 0:
                issues.append({
                    'type': '字段值一致性',
                    'error': f'存在{has_name_no_code}条记录有地类名称但无地类编码'
                })

        # 检查权属单位代码与名称的一致性
        if 'QSDWDM' in gdf.columns and 'QSDWMC' in gdf.columns:
            has_code_no_name = ((gdf['QSDWDM'].notna()) & (gdf['QSDWMC'].isna())).sum()
            if has_code_no_name > 0:
                issues.append({
                    'type': '字段值一致性',
                    'error': f'存在{has_code_no_name}条记录有权属单位代码但无单位名称'
                })

        # 检查坐落单位代码与名称的一致性
        if 'ZLDWDM' in gdf.columns and 'ZLDWMC' in gdf.columns:
            has_code_no_name = ((gdf['ZLDWDM'].notna()) & (gdf['ZLDWMC'].isna())).sum()
            if has_code_no_name > 0:
                issues.append({
                    'type': '字段值一致性',
                    'error': f'存在{has_code_no_name}条记录有坐落单位代码但无单位名称'
                })

    return issues

class GeoDataInspector:
    """地理数据质检器"""

    def __init__(self, input_dir: str, output_dir: Optional[str] = None, field_config_manager=None,
                 auto_fix_geometry=False, geometry_tolerance=0.001):
        """
        初始化检查器

        Args:
            input_dir: 输入目录路径
            output_dir: 输出目录路径，默认为当前目录
            field_config_manager: 字段配置管理器
            auto_fix_geometry: 是否自动修复几何错误
            geometry_tolerance: 几何修复容差
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir) if output_dir else Path.cwd()
        self.output_dir.mkdir(exist_ok=True)

        # 几何修复配置
        self.auto_fix_geometry = auto_fix_geometry
        self.geometry_tolerance = geometry_tolerance

        # 使用配置管理器中的字段标准
        if field_config_manager:
            self.field_standards = field_config_manager.get_field_standards()
        else:
            self.field_standards = FIELD_STANDARDS

        # 检查结果存储
        self.results = {
            'summary': {},
            'files': [],
            'errors': [],
            'topology_issues': [],
            'attribute_issues': [],
            'basic_issues': [],
            'geometry_fixes': []  # 新增：几何修复记录
        }

        # 存储所有几何数据和属性数据用于跨文件检查
        self.all_geometries = []
        self.all_dataframes = []
        self.file_indices = []

    def find_geospatial_files(self) -> List[Path]:
        """查找目录下的所有地理空间文件（SHP和GDB）"""
        geospatial_files = []

        # 查找SHP文件
        for file_path in self.input_dir.rglob("*.shp"):
            geospatial_files.append(file_path)

        # 查找GDB文件夹
        for gdb_path in self.input_dir.rglob("*.gdb"):
            if gdb_path.is_dir():
                geospatial_files.append(gdb_path)

        return geospatial_files

    def auto_fix_gaps_in_file(self, file_path: Path, tolerance: float = 0.001,
                             repair_method: str = 'buffer_merge') -> Dict:
        """
        自动修复文件中的面缝隙

        Args:
            file_path: 文件路径
            tolerance: 容差参数
            repair_method: 修复方法 ('buffer_merge', 'snap_vertices', 'extend_boundary')

        Returns:
            修复结果统计
        """
        try:
            logger.info(f"开始修复文件缝隙: {file_path.name}")

            # 使用改进的修复功能
            from improved_topology_utils import check_and_repair_gaps_in_file

            result = check_and_repair_gaps_in_file(
                str(file_path),
                tolerance=tolerance,
                repair_method=repair_method
            )

            if result['success']:
                logger.info(f"缝隙修复完成: 发现 {result['gaps_found']} 个缝隙, "
                          f"修复 {result.get('repaired_count', 0)} 个")
            else:
                logger.error(f"缝隙修复失败: {result.get('error', '未知错误')}")

            return result

        except ImportError:
            error_msg = "改进的拓扑修复模块不可用"
            logger.error(error_msg)
            return {'success': False, 'error': error_msg}
        except Exception as e:
            error_msg = f"缝隙修复失败: {str(e)}"
            logger.error(error_msg)
            return {'success': False, 'error': error_msg}

    def auto_fix_geometry_in_file(self, file_path: Path) -> Dict:
        """自动修复单个文件中的几何错误

        Args:
            file_path: 文件路径

        Returns:
            dict: 修复结果
        """
        try:
            logger.info(f"开始修复文件几何错误: {file_path.name}")

            # 读取文件
            try:
                gdf = gpd.read_file(file_path)
                if gdf.empty:
                    return {'success': False, 'error': '文件为空'}
            except Exception as e:
                logger.error(f"读取文件失败: {e}")
                return {'success': False, 'error': f'读取文件失败: {str(e)}'}

            # 转换几何类型以解决兼容性问题
            try:
                from improved_topology_utils import convert_geometry_types
                gdf = convert_geometry_types(gdf)
            except ImportError:
                logger.warning("几何类型转换模块不可用，跳过转换")
            except Exception as e:
                logger.warning(f"几何类型转换失败: {e}")

            # 统计修复前的几何问题
            total_geometries = len(gdf)
            invalid_before = sum(1 for geom in gdf.geometry if geom is not None and not geom.is_valid)

            if invalid_before == 0:
                return {'success': True, 'message': '没有发现几何错误', 'fixed_count': 0}

            # 修复几何错误
            fixed_count = 0
            error_count = 0
            fixed_indices = []

            for idx, row in gdf.iterrows():
                try:
                    geom = row.geometry
                    if geom is not None and not geom.is_valid:
                        # 尝试修复
                        fixed_geom = make_valid(geom)
                        if fixed_geom.is_valid:
                            gdf.at[idx, 'geometry'] = fixed_geom
                            fixed_count += 1
                            fixed_indices.append(idx)
                            logger.info(f"已修复几何 {idx}")
                        else:
                            error_count += 1
                            logger.warning(f"几何 {idx} 无法修复")
                except Exception as e:
                    error_count += 1
                    logger.error(f"修复几何 {idx} 时出错: {e}")

            # 保存修复后的文件
            if fixed_count > 0:
                # 创建备份文件
                backup_path = str(file_path) + '.backup'
                gdf.to_file(backup_path)
                logger.info(f"已创建备份文件: {backup_path}")

                # 保存修复后的文件
                gdf.to_file(file_path)
                logger.info(f"已保存修复后的文件: {file_path}")

                # 记录修复信息
                self.results['geometry_fixes'].append({
                    'file_path': str(file_path),
                    'file_name': file_path.name,
                    'total_geometries': total_geometries,
                    'invalid_before': invalid_before,
                    'fixed_count': fixed_count,
                    'error_count': error_count,
                    'fixed_indices': fixed_indices,
                    'backup_path': backup_path,
                    'fix_time': datetime.now().isoformat()
                })

            result = {
                'success': True,
                'total_geometries': total_geometries,
                'invalid_before': invalid_before,
                'fixed_count': fixed_count,
                'error_count': error_count,
                'fixed_indices': fixed_indices,
                'backup_path': backup_path if fixed_count > 0 else None
            }

            logger.info(f"几何修复完成: 修复 {fixed_count} 个，失败 {error_count} 个")
            return result

        except Exception as e:
            error_msg = f"自动修复几何文件失败: {str(e)}"
            logger.error(error_msg)
            return {'success': False, 'error': error_msg}

    def auto_fix_all_geometry_files(self, progress_callback=None) -> Dict:
        """自动修复所有文件的几何错误

        Args:
            progress_callback: 进度回调函数

        Returns:
            dict: 批量修复结果统计
        """
        try:
            logger.info("开始批量修复所有文件的几何错误")

            # 查找所有地理空间文件
            geospatial_files = self.find_geospatial_files()
            total_files = len(geospatial_files)

            if total_files == 0:
                return {'success': False, 'error': '没有找到地理空间文件'}

            successful_files = 0
            failed_files = 0
            total_fixed = 0
            total_errors = 0
            results = []

            for i, file_path in enumerate(geospatial_files):
                try:
                    if progress_callback:
                        progress_callback(i + 1, total_files, f"正在修复: {file_path.name}")

                    # 修复单个文件
                    result = self.auto_fix_geometry_in_file(file_path)
                    results.append({
                        'file_path': str(file_path),
                        'result': result
                    })

                    if result['success']:
                        successful_files += 1
                        if 'fixed_count' in result:
                            total_fixed += result['fixed_count']
                        if 'error_count' in result:
                            total_errors += result['error_count']
                    else:
                        failed_files += 1

                except Exception as e:
                    failed_files += 1
                    error_msg = f"处理文件 {file_path} 时出错: {str(e)}"
                    logger.error(error_msg)
                    results.append({
                        'file_path': str(file_path),
                        'result': {'success': False, 'error': error_msg}
                    })

            summary = {
                'success': True,
                'total_files': total_files,
                'successful_files': successful_files,
                'failed_files': failed_files,
                'total_fixed': total_fixed,
                'total_errors': total_errors,
                'results': results
            }

            logger.info(f"批量修复完成: 成功 {successful_files} 个文件，失败 {failed_files} 个文件，共修复 {total_fixed} 个几何错误")
            return summary

        except Exception as e:
            error_msg = f"批量修复几何文件失败: {str(e)}"
            logger.error(error_msg)
            return {'success': False, 'error': error_msg}

    def auto_fix_all_gaps_files(self, progress_callback=None) -> Dict:
        """自动修复所有文件的缝隙错误

        Args:
            progress_callback: 进度回调函数

        Returns:
            dict: 批量修复结果统计
        """
        try:
            logger.info("开始批量修复所有文件的缝隙错误")

            # 查找所有地理空间文件
            geospatial_files = self.find_geospatial_files()
            total_files = len(geospatial_files)

            if total_files == 0:
                return {'success': False, 'error': '没有找到地理空间文件'}

            successful_files = 0
            failed_files = 0
            total_gaps = 0
            repaired_gaps = 0
            results = []

            for i, file_path in enumerate(geospatial_files):
                try:
                    if progress_callback:
                        progress_callback(i + 1, total_files, f"正在修复缝隙: {file_path.name}")

                    # 修复单个文件的缝隙
                    result = self.auto_fix_gaps_in_file(file_path)
                    results.append({
                        'file_path': str(file_path),
                        'result': result
                    })

                    if result['success']:
                        successful_files += 1
                        if 'gaps_found' in result:
                            total_gaps += result['gaps_found']
                        if 'gaps_repaired' in result:
                            repaired_gaps += result['gaps_repaired']
                    else:
                        failed_files += 1

                except Exception as e:
                    logger.error(f"修复文件缝隙失败 {file_path}: {e}")
                    failed_files += 1
                    results.append({
                        'file_path': str(file_path),
                        'result': {'success': False, 'error': str(e)}
                    })

            summary = {
                'success': True,
                'total_files': total_files,
                'processed_files': successful_files,
                'failed_files': failed_files,
                'total_gaps': total_gaps,
                'repaired_gaps': repaired_gaps,
                'results': results
            }

            logger.info(f"批量缝隙修复完成: 成功 {successful_files} 个文件，失败 {failed_files} 个文件，共修复 {repaired_gaps} 个缝隙")
            return summary

        except Exception as e:
            error_msg = f"批量修复缝隙文件失败: {str(e)}"
            logger.error(error_msg)
            return {'success': False, 'error': error_msg}

    def comprehensive_fix_all_files(self, progress_callback=None) -> Dict:
        """综合修复所有文件：几何错误和缝隙错误

        Args:
            progress_callback: 进度回调函数

        Returns:
            dict: 综合修复结果统计
        """
        try:
            logger.info("开始综合修复所有文件")

            # 查找所有地理空间文件
            geospatial_files = self.find_geospatial_files()
            total_files = len(geospatial_files)

            if total_files == 0:
                return {'success': False, 'error': '没有找到地理空间文件'}

            # 第一步：修复几何错误
            if progress_callback:
                progress_callback(0, total_files * 2, "开始修复几何错误...")

            geometry_results = self.auto_fix_all_geometry_files(
                progress_callback=lambda current, total, msg:
                    progress_callback(current, total * 2, f"几何修复: {msg}") if progress_callback else None
            )

            # 第二步：修复缝隙错误
            if progress_callback:
                progress_callback(total_files, total_files * 2, "开始修复缝隙错误...")

            gap_results = self.auto_fix_all_gaps_files(
                progress_callback=lambda current, total, msg:
                    progress_callback(current + total_files, total_files * 2, f"缝隙修复: {msg}") if progress_callback else None
            )

            # 合并结果
            comprehensive_summary = {
                'success': True,
                'total_files': total_files,
                'processed_files': geometry_results.get('successful_files', 0),
                'geometry_results': geometry_results,
                'gap_results': gap_results,
                'backup_created': True  # 假设已创建备份
            }

            logger.info("综合修复完成")
            return comprehensive_summary

        except Exception as e:
            error_msg = f"综合修复文件失败: {str(e)}"
            logger.error(error_msg)
            return {'success': False, 'error': error_msg}

    def check_shp_file(self, shp_path: Path) -> Dict:
        """检查单个SHP文件的字段信息（优化大文件处理）"""
        # 记录检查开始时间
        check_start_time = datetime.now()

        result = {
            'file_name': shp_path.name,
            'file_path': str(shp_path),
            'geometry_type': None,
            'feature_count': 0,
            'fields': [],
            'field_count': 0,
            'file_size': 0,
            'related_files': [],
            'error': None,
            'topology_issues': [],
            'attribute_issues': [],
            'basic_issues': [],
            'file_hash': None,
            'check_start_time': check_start_time.isoformat(),
            'check_end_time': None
        }

        try:
            # 获取文件大小
            result['file_size'] = shp_path.stat().st_size

            # 计算文件哈希值
            result['file_hash'] = calculate_file_hash(shp_path)

            # 检查相关文件
            base_name = shp_path.stem
            parent_dir = shp_path.parent

            # 查找相关文件
            related_extensions = ['.dbf', '.shx', '.prj', '.cpg', '.qpj']
            for ext in related_extensions:
                related_file = parent_dir / f"{base_name}{ext}"
                if related_file.exists():
                    result['related_files'].append({
                        'extension': ext,
                        'file_name': related_file.name,
                        'file_size': related_file.stat().st_size,
                        'file_hash': calculate_file_hash(related_file),
                        'exists': True
                    })
                else:
                    result['related_files'].append({
                        'extension': ext,
                        'file_name': f"{base_name}{ext}",
                        'file_size': 0,
                        'file_hash': None,
                        'exists': False
                    })

            # 大文件处理优化
            file_size_mb = result['file_size'] / (1024 * 1024)
            chunk_size = 10000  # 默认分块大小

            if file_size_mb > 100:  # 大于100MB的文件
                chunk_size = 5000
                logger.info(f"大文件检测: {shp_path.name} ({file_size_mb:.1f}MB)，使用分块处理")
            elif file_size_mb > 50:  # 大于50MB的文件
                chunk_size = 8000
                logger.info(f"中等文件检测: {shp_path.name} ({file_size_mb:.1f}MB)，使用分块处理")

            # 读取SHP文件
            gdf = gpd.read_file(shp_path)

            # 获取几何类型
            if not gdf.empty:
                result['geometry_type'] = str(gdf.geometry.geom_type.iloc[0])
                result['feature_count'] = len(gdf)

                # 存储几何数据和属性数据用于跨文件检查（大文件只存储部分）
                if len(gdf) > chunk_size:
                    # 大文件只存储采样数据
                    sample_indices = np.linspace(0, len(gdf)-1, min(chunk_size, 10000), dtype=int)
                    sample_gdf = gdf.iloc[sample_indices]
                    self.all_geometries.extend(sample_gdf.geometry.tolist())
                    self.all_dataframes.append(sample_gdf)
                    self.file_indices.extend([len(self.all_dataframes)-1] * len(sample_gdf))
                    logger.info(f"大文件采样: {shp_path.name} 从 {len(gdf)} 个要素中采样 {len(sample_gdf)} 个")
                else:
                    self.all_geometries.extend(gdf.geometry.tolist())
                    self.all_dataframes.append(gdf)
                    self.file_indices.extend([len(self.all_dataframes)-1] * len(gdf))

                # 一、拓扑检测（大文件优化）
                if result['geometry_type'] in ['Polygon', 'MultiPolygon']:
                    # 对于大文件，只检查部分几何体
                    if len(gdf) > chunk_size:
                        sample_geometries = gdf.geometry.iloc[:chunk_size].tolist()
                        logger.info(f"大文件拓扑检查: {shp_path.name} 只检查前 {chunk_size} 个几何体")
                    else:
                        sample_geometries = gdf.geometry.tolist()

                    # 1. 面缝隙检测
                    gaps = check_topology_gaps(sample_geometries)
                    if gaps:
                        result['topology_issues'].extend(gaps)
                        self.results['topology_issues'].extend([{
                            'file': str(shp_path),
                            'issue': gap
                        } for gap in gaps])

                    # 2. 面重叠检测
                    overlaps = check_topology_overlaps(sample_geometries)
                    if overlaps:
                        result['topology_issues'].extend(overlaps)
                        self.results['topology_issues'].extend([{
                            'file': str(shp_path),
                            'issue': overlap
                        } for overlap in overlaps])

                # 二、属性检测（大文件优化）
                # 3. 数值范围符合性检查
                if len(gdf) > chunk_size:
                    sample_gdf = gdf.head(chunk_size)
                    range_issues = check_numeric_ranges(sample_gdf)
                    logger.info(f"大文件数值检查: {shp_path.name} 只检查前 {chunk_size} 行")
                else:
                    range_issues = check_numeric_ranges(gdf)

                if range_issues:
                    result['attribute_issues'].extend(range_issues)
                    self.results['attribute_issues'].extend([{
                        'file': str(shp_path),
                        'issue': issue
                    } for issue in range_issues])

                # 三、基础检测
                # 几何检查（大文件只检查部分）
                if len(gdf) > chunk_size:
                    sample_geometries = gdf.geometry.head(chunk_size).tolist()
                    sample_gdf = gdf.head(chunk_size)
                else:
                    sample_geometries = gdf.geometry.tolist()
                    sample_gdf = gdf

                # 几何检查（支持自动修复）
                if self.auto_fix_geometry:
                    # 自动修复模式
                    geom_issues, fixed_geometries = check_geometry_validity(sample_geometries, auto_fix=True, tolerance=self.geometry_tolerance)

                    # 如果有修复的几何，更新原始数据
                    if fixed_geometries:
                        logger.info(f"已自动修复 {len(fixed_geometries)} 个几何错误")
                        # 更新原始几何数据
                        for idx in fixed_geometries:
                            if idx < len(gdf):
                                gdf.at[idx, 'geometry'] = sample_geometries[idx]

                        # 保存修复后的文件
                        backup_path = str(shp_path) + '.backup'
                        gdf.to_file(backup_path)
                        logger.info(f"已创建备份文件: {backup_path}")
                        gdf.to_file(shp_path)
                        logger.info(f"已保存修复后的文件: {shp_path}")

                        # 记录修复信息
                        result['geometry_fixes'] = {
                            'fixed_count': len(fixed_geometries),
                            'backup_path': backup_path,
                            'fix_time': datetime.now().isoformat()
                        }
                else:
                    # 检测模式
                    geom_issues, _ = check_geometry_validity(sample_geometries, auto_fix=False)

                if geom_issues:
                    result['basic_issues'].extend(geom_issues)
                    self.results['basic_issues'].extend([{
                        'file': str(shp_path),
                        'issue': issue
                    } for issue in geom_issues])

                # 四、新增检查标准
                # 1. 数据完整性检查
                integrity_issues = check_data_integrity(sample_gdf)
                if integrity_issues:
                    result['basic_issues'].extend(integrity_issues)
                    self.results['basic_issues'].extend([{
                        'file': str(shp_path),
                        'issue': issue
                    } for issue in integrity_issues])

                # 2. 逻辑一致性检查
                logic_issues = check_logical_consistency(sample_gdf)
                if logic_issues:
                    result['attribute_issues'].extend(logic_issues)
                    self.results['attribute_issues'].extend([{
                        'file': str(shp_path),
                        'issue': issue
                    } for issue in logic_issues])

                # 3. 空间参考一致性检查
                spatial_issues = check_spatial_reference_consistency(sample_gdf)
                if spatial_issues:
                    result['basic_issues'].extend(spatial_issues)
                    self.results['basic_issues'].extend([{
                        'file': str(shp_path),
                        'issue': issue
                    } for issue in spatial_issues])

                # 4. 字段值一致性检查
                value_consistency_issues = check_field_value_consistency(sample_gdf)
                if value_consistency_issues:
                    result['attribute_issues'].extend(value_consistency_issues)
                    self.results['attribute_issues'].extend([{
                        'file': str(shp_path),
                        'issue': issue
                    } for issue in value_consistency_issues])

                # 5. 详细必填字段检查
                required_field_issues = check_required_fields_detailed(sample_gdf, self.field_standards, shp_path.name)
                if required_field_issues:
                    result['attribute_issues'].extend(required_field_issues)
                    self.results['attribute_issues'].extend([{
                        'file': str(shp_path),
                        'issue': issue
                    } for issue in required_field_issues])

            # 获取字段信息（大文件优化）
            if hasattr(gdf, 'columns'):
                # 对于大文件，只检查部分数据来评估字段
                sample_size = min(1000, len(gdf))
                sample_gdf = gdf.head(sample_size)

                for col in gdf.columns:
                    if col != 'geometry':
                        field_info = {
                            'name': col,
                            'type': str(gdf[col].dtype),
                            'null_count': sample_gdf[col].isnull().sum(),
                            'unique_count': sample_gdf[col].nunique(),
                            'sample_values': sample_gdf[col].dropna().head(3).tolist() if sample_gdf[col].dtype == 'object' else []
                        }
                        # 字段合规性检查 - 只检测表中实际存在的字段
                        if col in self.field_standards:
                            issues = check_field_compliance(col, sample_gdf[col], self.field_standards[col])
                            field_info['compliance_issues'] = issues
                        result['fields'].append(field_info)

                result['field_count'] = len(result['fields'])

        except Exception as e:
            result['error'] = str(e)
            logger.error(f"检查SHP文件失败 {shp_path}: {e}")

            # 即使有错误，也尝试收集字段信息
            try:
                if hasattr(gdf, 'columns'):
                    for col in gdf.columns:
                        if col != 'geometry':
                            field_info = {
                                'name': col,
                                'type': str(gdf[col].dtype),
                                'null_count': gdf[col].isnull().sum(),
                                'unique_count': gdf[col].nunique(),
                                'sample_values': gdf[col].dropna().head(3).tolist() if gdf[col].dtype == 'object' else []
                            }
                            # 字段合规性检查
                            if col in self.field_standards:
                                issues = check_field_compliance(col, gdf[col], self.field_standards[col])
                                field_info['compliance_issues'] = issues
                            result['fields'].append(field_info)

                    result['field_count'] = len(result['fields'])
            except Exception as field_error:
                logger.error(f"收集字段信息失败: {field_error}")

            # 标记几何错误为可忽略
            if "LinearRing" in str(e) or "linestring" in str(e).lower():
                # 将几何错误添加到几何问题中，而不是文件读取错误
                geometry_issue = f"几何错误 - 可忽略: {str(e)}"
                result['topology_issues'].append(geometry_issue)
                self.results['topology_issues'].append({
                    'file': str(shp_path),
                    'issue': geometry_issue
                })
            else:
                self.results['errors'].append({
                    'file': str(shp_path),
                    'error': str(e),
                    'level': ERROR_LEVELS['CRITICAL'],
                    'type': ERROR_TYPES['OTHER_ERROR']
                })
        finally:
            # 记录检查结束时间
            result['check_end_time'] = datetime.now().isoformat()

        return result

    def check_dbf_file(self, dbf_path: Path) -> Dict:
        """检查单个DBF文件的字段信息"""
        # 记录检查开始时间
        check_start_time = datetime.now()

        result = {
            'file_name': dbf_path.name,
            'file_path': str(dbf_path),
            'fields': [],
            'field_count': 0,
            'file_size': 0,
            'error': None,
            'file_hash': None,
            'check_start_time': check_start_time.isoformat(),
            'check_end_time': None
        }

        try:
            # 获取文件大小
            result['file_size'] = dbf_path.stat().st_size

            # 计算文件哈希值
            result['file_hash'] = calculate_file_hash(dbf_path)

            # 尝试多种编码读取DBF文件
            df = None
            encodings = ['gbk', 'gb2312', 'utf-8', 'latin1']

            for encoding in encodings:
                try:
                    df = pd.read_csv(dbf_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                # 如果所有编码都失败，尝试使用二进制模式读取基本信息
                result['error'] = "无法使用任何编码读取DBF文件"
                self.results['errors'].append({
                    'file': str(dbf_path),
                    'error': "编码错误 - 可忽略"
                })
                return result

            # 获取字段信息
            for col in df.columns:
                field_info = {
                    'name': col,
                    'type': str(df[col].dtype),
                    'null_count': df[col].isnull().sum(),
                    'unique_count': df[col].nunique(),
                    'sample_values': df[col].dropna().head(3).tolist() if df[col].dtype == 'object' else []
                }
                # 字段合规性检查
                if col in self.field_standards:
                    issues = check_field_compliance(col, df[col], self.field_standards[col])
                    field_info['compliance_issues'] = issues
                result['fields'].append(field_info)

            result['field_count'] = len(result['fields'])

        except Exception as e:
            result['error'] = str(e)
            # 标记为可忽略的错误
            if "codec can't decode" in str(e):
                self.results['errors'].append({
                    'file': str(dbf_path),
                    'error': f"编码错误 - 可忽略: {str(e)}",
                    'level': ERROR_LEVELS['IGNORABLE'],
                    'type': ERROR_TYPES['ENCODING_ERROR']
                })
            else:
                self.results['errors'].append({
                    'file': str(dbf_path),
                    'error': str(e),
                    'level': ERROR_LEVELS['CRITICAL'],
                    'type': ERROR_TYPES['OTHER_ERROR']
                })
        finally:
            # 记录检查结束时间
            result['check_end_time'] = datetime.now().isoformat()

        return result

    def check_gdb_file(self, gdb_path: Path) -> Dict:
        """检查单个GDB文件的字段信息"""
        # 记录检查开始时间
        check_start_time = datetime.now()

        result = {
            'file_name': gdb_path.name,
            'file_path': str(gdb_path),
            'file_type': 'GDB',
            'geometry_type': None,
            'feature_count': 0,
            'fields': [],
            'field_count': 0,
            'file_size': 0,
            'layers': [],
            'layer_count': 0,
            'error': None,
            'file_hash': None,
            'check_start_time': check_start_time.isoformat(),
            'check_end_time': None,
            'topology_issues': [],
            'attribute_issues': [],
            'basic_issues': []
        }

        try:
            # 获取文件夹大小（GDB是文件夹）
            if gdb_path.is_dir():
                total_size = 0
                for file_path in gdb_path.rglob('*'):
                    if file_path.is_file():
                        total_size += file_path.stat().st_size
                result['file_size'] = total_size
            else:
                result['file_size'] = gdb_path.stat().st_size

            # 计算文件夹哈希值（简化版本，只计算主要文件）
            result['file_hash'] = self._calculate_gdb_hash(gdb_path)

            # 读取GDB中的所有图层
            try:
                layers = gpd.read_file(gdb_path, driver='OpenFileGDB')
                if isinstance(layers, gpd.GeoDataFrame):
                    # 单个图层
                    layers = [layers]

                result['layer_count'] = len(layers)
                total_features = 0
                all_fields = set()

                for i, layer in enumerate(layers):
                    layer_info = {
                        'layer_name': f'Layer_{i+1}',
                        'feature_count': len(layer),
                        'geometry_type': str(layer.geometry.geom_type.iloc[0]) if len(layer) > 0 else 'Unknown',
                        'fields': []
                    }

                    # 获取字段信息
                    for col in layer.columns:
                        if col != 'geometry':
                            field_info = {
                                'name': col,
                                'type': str(layer[col].dtype),
                                'null_count': layer[col].isnull().sum(),
                                'unique_count': layer[col].nunique(),
                                'sample_values': layer[col].dropna().head(3).tolist() if layer[col].dtype == 'object' else []
                            }

                            # 字段合规性检查
                            if col in self.field_standards:
                                issues = check_field_compliance(col, layer[col], self.field_standards[col])
                                field_info['compliance_issues'] = issues

                            layer_info['fields'].append(field_info)
                            all_fields.add(col)

                    layer_info['field_count'] = len(layer_info['fields'])
                    result['layers'].append(layer_info)
                    total_features += layer_info['feature_count']

                    # 存储数据用于跨文件检查
                    self.all_dataframes.append(layer)
                    self.all_geometries.extend(layer.geometry.tolist())
                    self.file_indices.append(len(self.all_dataframes) - 1)

                result['feature_count'] = total_features
                result['field_count'] = len(all_fields)
                result['geometry_type'] = 'Multiple' if len(layers) > 1 else (layers[0].geometry.geom_type.iloc[0] if len(layers) > 0 else 'Unknown')

                # 几何检查
                if self.all_geometries:
                    # 拓扑检查
                    topology_gaps = check_topology_gaps(self.all_geometries)
                    topology_overlaps = check_topology_overlaps(self.all_geometries)
                    geometry_validity = check_geometry_validity(self.all_geometries)

                    if topology_gaps:
                        result['topology_issues'].extend(topology_gaps)
                    if topology_overlaps:
                        result['topology_issues'].extend(topology_overlaps)
                    if geometry_validity:
                        result['basic_issues'].extend(geometry_validity)

                    # 坐标系统检查
                    if layers:
                        crs_issues = check_coordinate_system(layers[0])
                        if crs_issues:
                            result['basic_issues'].extend(crs_issues)

                # 新增检查标准 - 对每个图层进行检查
                for i, layer in enumerate(layers):
                    # 1. 数据完整性检查
                    integrity_issues = check_data_integrity(layer)
                    if integrity_issues:
                        result['basic_issues'].extend(integrity_issues)

                    # 2. 逻辑一致性检查
                    logic_issues = check_logical_consistency(layer)
                    if logic_issues:
                        result['attribute_issues'].extend(logic_issues)

                    # 3. 空间参考一致性检查
                    spatial_issues = check_spatial_reference_consistency(layer)
                    if spatial_issues:
                        result['basic_issues'].extend(spatial_issues)

                    # 4. 字段值一致性检查
                    value_consistency_issues = check_field_value_consistency(layer)
                    if value_consistency_issues:
                        result['attribute_issues'].extend(value_consistency_issues)

                    # 5. 详细必填字段检查
                    required_field_issues = check_required_fields_detailed(layer, self.field_standards, gdb_path.name)
                    if required_field_issues:
                        result['attribute_issues'].extend(required_field_issues)

            except Exception as e:
                result['error'] = f"无法读取GDB文件: {str(e)}"
                self.results['errors'].append({
                    'file': str(gdb_path),
                    'error': f"GDB读取错误: {str(e)}",
                    'level': ERROR_LEVELS['CRITICAL'],
                    'type': ERROR_TYPES['OTHER_ERROR']
                })

        except Exception as e:
            result['error'] = str(e)
            self.results['errors'].append({
                'file': str(gdb_path),
                'error': str(e),
                'level': ERROR_LEVELS['CRITICAL'],
                'type': ERROR_TYPES['OTHER_ERROR']
            })
        finally:
            # 记录检查结束时间
            result['check_end_time'] = datetime.now().isoformat()

        return result

    def _calculate_gdb_hash(self, gdb_path: Path) -> str:
        """计算GDB文件夹的哈希值（简化版本）"""
        try:
            hash_obj = hashlib.sha256()

            # 只计算主要文件
            main_files = ['gdb', 'freelist', 'a00000001.gdbtable', 'a00000001.gdbindexes']

            for file_name in main_files:
                file_path = gdb_path / file_name
                if file_path.exists():
                    with open(file_path, 'rb') as f:
                        for chunk in iter(lambda: f.read(4096), b""):
                            hash_obj.update(chunk)

            return hash_obj.hexdigest()
        except Exception:
            return "无法计算哈希值"

    def run_check(self, progress_callback=None) -> Dict:
        """运行检查"""
        logger.info(f"开始检查目录: {self.input_dir}")

        # 查找地理空间文件
        geospatial_files = self.find_geospatial_files()
        logger.info(f"找到 {len(geospatial_files)} 个地理空间文件")

        # 查找DBF文件
        dbf_files = []
        for file_path in self.input_dir.rglob("*.dbf"):
            dbf_files.append(file_path)
        logger.info(f"找到 {len(dbf_files)} 个DBF文件")

        total_files = len(geospatial_files) + len(dbf_files)
        processed_files = 0

        # 检查地理空间文件
        for i, file_path in enumerate(geospatial_files):
            logger.info(f"正在检查 ({i+1}/{len(geospatial_files)}): {file_path.name}")
            if file_path.suffix.lower() == '.shp':
                result = self.check_shp_file(file_path)
            elif file_path.suffix.lower() == '.gdb' or file_path.is_dir():
                result = self.check_gdb_file(file_path)
            else:
                continue
            self.results['files'].append(result)
            processed_files += 1
            if progress_callback:
                progress_callback(processed_files, total_files, f"检查文件: {file_path.name}")

        # 检查DBF文件
        for i, dbf_path in enumerate(dbf_files):
            logger.info(f"正在检查DBF ({i+1}/{len(dbf_files)}): {dbf_path.name}")
            result = self.check_dbf_file(dbf_path)
            self.results['files'].append(result)
            processed_files += 1
            if progress_callback:
                progress_callback(processed_files, total_files, f"检查DBF文件: {dbf_path.name}")

        # 进行跨文件检查
        if progress_callback:
            progress_callback(processed_files, total_files, "进行跨文件检查...")

        # 二、属性检测
        # 1. 属性结构一致性检查
        if self.all_dataframes:
            structure_issues = check_attribute_structure_consistency(self.all_dataframes)
            if structure_issues:
                self.results['attribute_issues'].extend([{
                    'file': '跨文件检查',
                    'issue': issue
                } for issue in structure_issues])

        # 2. 编号唯一性检查
        if self.all_dataframes:
            uniqueness_issues = check_unique_identifiers(self.all_dataframes)
            if uniqueness_issues:
                self.results['attribute_issues'].extend([{
                    'file': '跨文件检查',
                    'issue': issue
                } for issue in uniqueness_issues])

        # 生成摘要
        self._generate_summary()

        return self.results

    def _generate_summary(self):
        """生成检查结果摘要"""
        total_files = len(self.results['files'])
        shp_files = [f for f in self.results['files'] if f['file_name'].endswith('.shp')]
        gdb_files = [f for f in self.results['files'] if f.get('file_type') == 'GDB']
        dbf_files = [f for f in self.results['files'] if f['file_name'].endswith('.dbf')]
        error_files = len(self.results['errors'])

        # 计算总要素数量（包括GDB文件）
        total_features = sum(f.get('feature_count', 0) for f in shp_files + gdb_files)
        total_fields = sum(f.get('field_count', 0) for f in self.results['files'])

        # 统计新检查结果
        topology_issues = len(self.results['topology_issues'])
        attribute_issues = len(self.results['attribute_issues'])
        basic_issues = len(self.results['basic_issues'])

        self.results['summary'] = {
            'check_time': datetime.now().isoformat(),
            'total_files': total_files,
            'shp_files': len(shp_files),
            'gdb_files': len(gdb_files),
            'dbf_files': len(dbf_files),
            'error_files': error_files,
            'total_features': total_features,
            'total_fields': total_fields,
            'topology_issues': topology_issues,
            'attribute_issues': attribute_issues,
            'basic_issues': basic_issues
        }

    def _to_serializable(self, obj):
        """将对象转换为可序列化的格式"""
        if isinstance(obj, np.integer):
            return int(obj)
        elif isinstance(obj, np.floating):
            return float(obj)
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        return obj

    def save_results(self, format: str = 'json') -> str:
        """保存检查结果"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if format == 'json':
            filename = f"shp_field_check_results_{timestamp}.json"
            filepath = self.output_dir / filename

            # 转换不可序列化的对象
            serializable_results = json.loads(
                json.dumps(self.results, default=self._to_serializable)
            )

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(serializable_results, f, ensure_ascii=False, indent=2)

        elif format == 'excel':
            filename = f"shp_field_check_results_{timestamp}.xlsx"
            filepath = self.output_dir / filename

            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # 摘要信息
                summary_df = pd.DataFrame([self.results['summary']])
                summary_df.to_excel(writer, sheet_name='摘要', index=False)

                # 文件检查结果
                files_data = []
                for file_result in self.results['files']:
                    files_data.append({
                        '文件名': file_result['file_name'],
                        '文件路径': file_result['file_path'],
                        '几何类型': file_result.get('geometry_type', ''),
                        '要素数量': file_result.get('feature_count', 0),
                        '字段数量': file_result.get('field_count', 0),
                        '文件大小(KB)': round(file_result.get('file_size', 0) / 1024, 2),
                        '文件哈希值': file_result.get('file_hash', ''),
                        '检查开始时间': file_result.get('check_start_time', ''),
                        '检查结束时间': file_result.get('check_end_time', ''),
                        '错误信息': file_result.get('error', '')
                    })

                if files_data:
                    files_df = pd.DataFrame(files_data)
                    files_df.to_excel(writer, sheet_name='文件检查结果', index=False)

                # 字段合规性详情
                compliance_data = []
                for file_result in self.results['files']:
                    for field in file_result.get('fields', []):
                        compliance_data.append({
                            '文件名': file_result['file_name'],
                            '字段名': field['name'],
                            '字段类型': field['type'],
                            '空值数量': field['null_count'],
                            '唯一值数量': field['unique_count'],
                            '合规性问题': '; '.join(field.get('compliance_issues', []))
                        })

                if compliance_data:
                    compliance_df = pd.DataFrame(compliance_data)
                    compliance_df.to_excel(writer, sheet_name='字段合规性', index=False)

        elif format == 'word':
            filename = f"地理数据质检报告_{timestamp}.docx"
            filepath = self.output_dir / filename
            self._generate_word_report(filepath)

        return str(filepath)

    def _generate_word_report(self, filepath: Path):
        """生成Word格式的正式报告"""
        doc = Document()

        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # 标题
        title = doc.add_heading('地理数据质检报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 报告信息
        doc.add_paragraph()
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'

        info_table.cell(0, 0).text = '检查时间'
        info_table.cell(0, 1).text = self.results['summary']['check_time']
        info_table.cell(1, 0).text = '检查目录'
        info_table.cell(1, 1).text = str(self.input_dir)
        info_table.cell(2, 0).text = '输出目录'
        info_table.cell(2, 1).text = str(self.output_dir)
        info_table.cell(3, 0).text = '报告生成时间'
        info_table.cell(3, 1).text = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
        info_table.cell(4, 0).text = '哈希算法'
        info_table.cell(4, 1).text = 'SHA256'

        # 检查摘要
        doc.add_heading('1. 检查摘要', level=1)
        summary = self.results['summary']

        summary_text = f"""
本次检查共处理文件 {summary['total_files']} 个，其中SHP文件 {summary['shp_files']} 个，DBF文件 {summary['dbf_files']} 个。
总要素数量：{summary['total_features']} 个
总字段数量：{summary['total_fields']} 个
错误文件数量：{summary['error_files']} 个
        """
        doc.add_paragraph(summary_text.strip())

        # 文件详细信息
        doc.add_heading('2. 文件详细信息', level=1)

        # 创建文件详细信息表格
        files_table = doc.add_table(rows=1, cols=6)
        files_table.style = 'Table Grid'
        files_table.cell(0, 0).text = '文件名'
        files_table.cell(0, 1).text = '文件大小(KB)'
        files_table.cell(0, 2).text = '文件哈希值'
        files_table.cell(0, 3).text = '检查开始时间'
        files_table.cell(0, 4).text = '检查结束时间'
        files_table.cell(0, 5).text = '状态'

        # 添加文件信息行
        for file_result in self.results['files']:
            row = files_table.add_row()
            row.cells[0].text = file_result['file_name']
            row.cells[1].text = f"{round(file_result.get('file_size', 0) / 1024, 2)}"
            row.cells[2].text = file_result.get('file_hash', '')[:16] + '...' if file_result.get('file_hash') else '计算失败'
            row.cells[3].text = file_result.get('check_start_time', '')[:19] if file_result.get('check_start_time') else ''
            row.cells[4].text = file_result.get('check_end_time', '')[:19] if file_result.get('check_end_time') else ''
            row.cells[5].text = '正常' if not file_result.get('error') else '错误'

        doc.add_paragraph()

        # 错误分类统计
        doc.add_heading('3. 错误分类统计', level=1)

        ignorable_errors = [e for e in (self.results.get('errors', []) if self.results else []) if '可忽略' in e.get('error', '')]
        critical_errors = [e for e in (self.results.get('errors', []) if self.results else []) if '可忽略' not in e.get('error', '')]

        error_stats = doc.add_table(rows=3, cols=3)
        error_stats.style = 'Table Grid'
        error_stats.cell(0, 0).text = '错误类型'
        error_stats.cell(0, 1).text = '数量'
        error_stats.cell(0, 2).text = '严重程度'

        error_stats.cell(1, 0).text = '可忽略错误'
        error_stats.cell(1, 1).text = str(len(ignorable_errors))
        error_stats.cell(1, 2).text = ERROR_LEVELS['IGNORABLE']

        error_stats.cell(2, 0).text = '不可忽略错误'
        error_stats.cell(2, 1).text = str(len(critical_errors))
        error_stats.cell(2, 2).text = ERROR_LEVELS['CRITICAL']

        # 详细错误信息
        if self.results and self.results.get('errors'):
            doc.add_heading('4. 详细错误信息', level=1)

            # 可忽略错误
            if ignorable_errors:
                doc.add_heading('4.1 可忽略错误', level=2)
                ignorable_table = doc.add_table(rows=1, cols=3)
                ignorable_table.style = 'Table Grid'
                ignorable_table.cell(0, 0).text = '文件名'
                ignorable_table.cell(0, 1).text = '错误类型'
                ignorable_table.cell(0, 2).text = '错误描述'

                for error in ignorable_errors:
                    row = ignorable_table.add_row()
                    row.cells[0].text = Path(error['file']).name
                    row.cells[1].text = ERROR_TYPES['ENCODING_ERROR'] if '编码' in error['error'] else ERROR_TYPES['GEOMETRY_ERROR']
                    row.cells[2].text = error['error']

            # 不可忽略错误
            if critical_errors:
                doc.add_heading('4.2 不可忽略错误', level=2)
                critical_table = doc.add_table(rows=1, cols=3)
                critical_table.style = 'Table Grid'
                critical_table.cell(0, 0).text = '文件名'
                critical_table.cell(0, 1).text = '错误类型'
                critical_table.cell(0, 2).text = '错误描述'

                for error in critical_errors:
                    row = critical_table.add_row()
                    row.cells[0].text = Path(error['file']).name
                    row.cells[1].text = ERROR_TYPES['OTHER_ERROR']
                    row.cells[2].text = error['error']

        # 字段合规性检查结果
        doc.add_heading('5. 字段合规性检查结果', level=1)

        # 统计合规性
        total_fields = 0
        compliant_fields = 0
        non_compliant_fields = 0

        for file_result in self.results['files']:
            for field in file_result.get('fields', []):
                total_fields += 1
                if field.get('compliance_issues'):
                    non_compliant_fields += 1
                else:
                    compliant_fields += 1

        # 计算合规率，避免除零错误
        compliance_rate = (compliant_fields/total_fields*100) if total_fields > 0 else 0.0

        compliance_text = f"""
字段合规性统计：
总字段数：{total_fields} 个
合规字段：{compliant_fields} 个
不合规字段：{non_compliant_fields} 个
合规率：{compliance_rate:.1f}%
        """
        doc.add_paragraph(compliance_text.strip())

        # 不合规字段详情
        if non_compliant_fields > 0:
            doc.add_heading('5.1 不合规字段详情', level=2)
            non_compliant_table = doc.add_table(rows=1, cols=4)
            non_compliant_table.style = 'Table Grid'
            non_compliant_table.cell(0, 0).text = '文件名'
            non_compliant_table.cell(0, 1).text = '字段名'
            non_compliant_table.cell(0, 2).text = '字段类型'
            non_compliant_table.cell(0, 3).text = '合规性问题'

            for file_result in self.results['files']:
                for field in file_result.get('fields', []):
                    if field.get('compliance_issues'):
                        row = non_compliant_table.add_row()
                        row.cells[0].text = file_result['file_name']
                        row.cells[1].text = field['name']
                        row.cells[2].text = field['type']
                        row.cells[3].text = '; '.join(field['compliance_issues'])

        # 检查结论
        doc.add_heading('6. 检查结论', level=1)

        # 计算合规率，避免除零错误
        compliance_rate = (compliant_fields/total_fields*100) if total_fields > 0 else 0.0

        conclusion = f"""
基于本次检查结果，得出以下结论：

1. 文件完整性：共检查 {summary['total_files']} 个文件，其中 {summary['error_files']} 个文件存在问题
2. 错误严重程度：
   - 可忽略错误：{len(ignorable_errors)} 个（主要为编码和几何问题）
   - 不可忽略错误：{len(critical_errors)} 个
3. 字段合规性：合规率 {compliance_rate:.1f}%

建议：
- 对于可忽略错误，建议在数据预处理阶段进行编码转换和几何修复
- 对于不可忽略错误，需要立即处理以确保数据质量
- 对于不合规字段，建议按照标准规范进行修正
        """
        doc.add_paragraph(conclusion.strip())

        # 保存文档
        doc.save(str(filepath))

    def print_summary(self):
        """打印检查结果摘要"""
        summary = self.results['summary']
        print("\n" + "="*50)
        print("检查结果摘要")
        print("="*50)
        print(f"检查时间: {summary['check_time']}")
        print(f"检查文件总数: {summary['total_files']}")
        print(f"SHP文件数量: {summary['shp_files']}")
        print(f"DBF文件数量: {summary['dbf_files']}")
        print(f"错误文件数量: {summary['error_files']}")
        print(f"总要素数量: {summary['total_features']}")
        print(f"总字段数量: {summary['total_fields']}")

        if summary['error_files'] > 0:
            print("\n错误文件列表:")
            for error in (self.results.get('errors', []) if self.results else []):
                print(f"  - {Path(error['file']).name}: {error['error']}")

    def print_detailed_results(self):
        """打印详细检查结果"""
        print("\n" + "="*50)
        print("详细检查结果")
        print("="*50)

        files = self.results.get('files', []) if self.results and isinstance(self.results, dict) else []
        for file_result in files:
            print(f"\n文件: {file_result['file_name']}")
            print(f"路径: {file_result['file_path']}")

            if file_result.get('error'):
                print(f"错误: {file_result['error']}")
                continue

            print(f"几何类型: {file_result.get('geometry_type', 'N/A')}")
            print(f"要素数量: {file_result.get('feature_count', 0)}")
            print(f"字段数量: {file_result.get('field_count', 0)}")
            print(f"文件大小: {file_result.get('file_size', 0)} bytes")
            print(f"文件哈希值: {file_result.get('file_hash', 'N/A')}")
            print(f"检查开始时间: {file_result.get('check_start_time', 'N/A')}")
            print(f"检查结束时间: {file_result.get('check_end_time', 'N/A')}")

            # 相关文件
            if file_result.get('related_files'):
                print("相关文件:")
                for related in file_result['related_files']:
                    status = "存在" if related['exists'] else "缺失"
                    hash_info = f"哈希值: {related.get('file_hash', 'N/A')}" if related['exists'] else "哈希值: N/A"
                    print(f"  {related['file_name']}: {status} ({hash_info})")

            # 字段信息
            if file_result.get('fields'):
                print("字段信息:")
                for field in file_result.get('fields', []):
                    print(f"  {field['name']} ({field['type']})")
                    print(f"    空值数量: {field['null_count']}")
                    print(f"    唯一值数量: {field['unique_count']}")

                    if field.get('compliance_issues'):
                        print(f"    合规性问题: {', '.join(field['compliance_issues'])}")

                    if field.get('sample_values'):
                        print(f"    样本值: {field['sample_values'][:3]}")

# FieldConfigDialog 类已移除，现在使用 PandasTable 版本

class GeoDataInspectorGUI:
    """地理数据质检器GUI界面"""

    def __init__(self):
        self.root = tk.Tk()
        self.root.title("地理数据质检工具 v2.0")
        self.root.geometry("1400x900")

        # 配置系统字体
        configure_system_fonts()

        # 设置窗口图标（如果有的话）
        try:
            icon_path = self.get_resource_path("favicon.ico")
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except:
            pass

        # 初始化字段配置管理器
        self.field_config_manager = FieldConfigManager()

        # 检查结果
        self.results = None
        self.checker = None

        # 记住上次的目录
        self.last_input_dir = ""
        self.last_output_dir = ""

        # 新增：进度相关变量
        self.start_time = None
        self.end_time = None
        self.current_file = ""
        self.current_phase = ""
        self.estimated_remaining = ""
        self.memory_usage = ""

        # 几何修复配置
        self.auto_fix_geometry_var = tk.BooleanVar(value=False)
        self.geometry_tolerance_var = tk.DoubleVar(value=0.001)

        self.setup_ui()
        self.load_last_directories()

    def get_resource_path(self, relative_path):
        """获取资源文件路径（支持PyInstaller打包）"""
        try:
            # PyInstaller创建的临时文件夹
            base_path = getattr(sys, '_MEIPASS', os.path.abspath("."))
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)



    def setup_ui(self):
        """设置用户界面"""
        # 主框架
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题区域
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=(0, 20))

        title_label = ttk.Label(title_frame, text="地理数据质检工具",
                               font=("TkDefaultFont", 18, "bold"), foreground="#2E86AB")
        title_label.pack(side=tk.LEFT)

        # 版本信息
        version_label = ttk.Label(title_frame, text="v2.1",
                                 font=("TkDefaultFont", 10), foreground="#666666")
        version_label.pack(side=tk.RIGHT, pady=(0, 5))

        # 输入输出区域
        io_frame = ttk.LabelFrame(main_frame, text="📁 输入输出设置", padding="15")
        io_frame.pack(fill=tk.X, pady=(0, 15))

        # 输入目录
        input_frame = ttk.Frame(io_frame)
        input_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(input_frame, text="📂 输入目录:", font=("TkDefaultFont", 10, "bold")).pack(side=tk.LEFT)
        self.input_dir_var = tk.StringVar()
        input_entry = ttk.Entry(input_frame, textvariable=self.input_dir_var, width=60, font=("TkDefaultFont", 9))
        input_entry.pack(side=tk.LEFT, padx=(10, 10), fill=tk.X, expand=True)
        ttk.Button(input_frame, text="浏览", command=self.browse_input_dir).pack(side=tk.RIGHT)

        # 输出目录
        output_frame = ttk.Frame(io_frame)
        output_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(output_frame, text="📁 输出目录:", font=("TkDefaultFont", 10, "bold")).pack(side=tk.LEFT)
        self.output_dir_var = tk.StringVar(value=str(Path.cwd()))
        output_entry = ttk.Entry(output_frame, textvariable=self.output_dir_var, width=60, font=("TkDefaultFont", 9))
        output_entry.pack(side=tk.LEFT, padx=(10, 10), fill=tk.X, expand=True)
        ttk.Button(output_frame, text="浏览", command=self.browse_output_dir).pack(side=tk.RIGHT)

        # 控制按钮区域
        control_frame = ttk.LabelFrame(main_frame, text="⚙️ 操作控制", padding="15")
        control_frame.pack(fill=tk.X, pady=(0, 15))

        # 左侧按钮
        left_buttons = ttk.Frame(control_frame)
        left_buttons.pack(side=tk.LEFT)

        self.check_button = ttk.Button(left_buttons, text="🚀 开始检查", command=self.start_check)
        self.check_button.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(left_buttons, text="⚙️ 字段配置", command=self.open_field_config).pack(side=tk.LEFT, padx=(0, 10))

        self.export_button = ttk.Button(left_buttons, text="📊 导出报告", command=self.export_report, state=tk.DISABLED)
        self.export_button.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(left_buttons, text="🗑️ 清空结果", command=self.clear_results).pack(side=tk.LEFT)

        # 几何修复选项
        geometry_frame = ttk.Frame(control_frame)
        geometry_frame.pack(side=tk.LEFT, padx=(20, 0))

        # 自动修复复选框
        auto_fix_check = ttk.Checkbutton(geometry_frame, text="🔧 自动修复几何",
                                        variable=self.auto_fix_geometry_var)
        auto_fix_check.pack(side=tk.LEFT, padx=(0, 10))

        # 容差设置
        ttk.Label(geometry_frame, text="容差:", font=("Arial", 9)).pack(side=tk.LEFT, padx=(0, 5))
        tolerance_entry = ttk.Entry(geometry_frame, textvariable=self.geometry_tolerance_var,
                                   width=8, font=("Arial", 9))
        tolerance_entry.pack(side=tk.LEFT, padx=(0, 10))

        # 几何修复按钮已移至批量修复区域，避免功能重复

        # 右侧状态显示
        right_frame = ttk.Frame(control_frame)
        right_frame.pack(side=tk.RIGHT)

        # 文件计数显示
        self.file_count_var = tk.StringVar(value="文件: 0")
        file_count_label = ttk.Label(right_frame, textvariable=self.file_count_var,
                                    font=("Arial", 9), foreground="#666666")
        file_count_label.pack(side=tk.RIGHT, padx=(10, 0))

        # 进度条区域
        progress_frame = ttk.LabelFrame(main_frame, text="📈 检查进度", padding="10")
        progress_frame.pack(fill=tk.X, pady=(0, 15))

        # 进度条和状态
        progress_inner_frame = ttk.Frame(progress_frame)
        progress_inner_frame.pack(fill=tk.X)

        # 进度条
        progress_bar_frame = ttk.Frame(progress_inner_frame)
        progress_bar_frame.pack(fill=tk.X, pady=(0, 5))

        ttk.Label(progress_bar_frame, text="进度:", font=("Arial", 9, "bold")).pack(side=tk.LEFT)
        self.progress_var = tk.DoubleVar()
        progress_bar = ttk.Progressbar(progress_bar_frame, variable=self.progress_var,
                                     maximum=100, length=400, mode='determinate')
        progress_bar.pack(side=tk.LEFT, padx=(10, 10), fill=tk.X, expand=True)

        # 状态标签
        self.status_var = tk.StringVar(value="就绪")
        status_label = ttk.Label(progress_bar_frame, textvariable=self.status_var,
                                font=("Arial", 9), foreground="#2E86AB")
        status_label.pack(side=tk.RIGHT)

        # 详细信息区域
        details_frame = ttk.Frame(progress_inner_frame)
        details_frame.pack(fill=tk.X)

        # 左侧信息
        left_info = ttk.Frame(details_frame)
        left_info.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # 当前文件
        current_file_frame = ttk.Frame(left_info)
        current_file_frame.pack(fill=tk.X, pady=2)
        ttk.Label(current_file_frame, text="当前文件:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.current_file_var = tk.StringVar(value="无")
        ttk.Label(current_file_frame, textvariable=self.current_file_var,
                 font=("Arial", 8), foreground="#666666").pack(side=tk.LEFT, padx=(5, 0))

        # 处理阶段
        phase_frame = ttk.Frame(left_info)
        phase_frame.pack(fill=tk.X, pady=2)
        ttk.Label(phase_frame, text="处理阶段:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.current_phase_var = tk.StringVar(value="无")
        ttk.Label(phase_frame, textvariable=self.current_phase_var,
                 font=("Arial", 8), foreground="#666666").pack(side=tk.LEFT, padx=(5, 0))

        # 右侧信息
        right_info = ttk.Frame(details_frame)
        right_info.pack(side=tk.RIGHT)

        # 预估剩余时间
        time_frame = ttk.Frame(right_info)
        time_frame.pack(fill=tk.X, pady=2)
        ttk.Label(time_frame, text="剩余时间:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.estimated_time_var = tk.StringVar(value="--")
        ttk.Label(time_frame, textvariable=self.estimated_time_var,
                 font=("Arial", 8), foreground="#666666").pack(side=tk.LEFT, padx=(5, 0))

        # 内存使用
        memory_frame = ttk.Frame(right_info)
        memory_frame.pack(fill=tk.X, pady=2)
        ttk.Label(memory_frame, text="内存使用:", font=("Arial", 8)).pack(side=tk.LEFT)
        self.memory_usage_var = tk.StringVar(value="--")
        ttk.Label(memory_frame, textvariable=self.memory_usage_var,
                 font=("Arial", 8), foreground="#666666").pack(side=tk.LEFT, padx=(5, 0))

        # 结果显示区域
        result_frame = ttk.LabelFrame(main_frame, text="📋 检查结果", padding="10")
        result_frame.pack(fill=tk.BOTH, expand=True)

        # 创建Notebook用于标签页
        notebook = ttk.Notebook(result_frame)
        notebook.pack(fill=tk.BOTH, expand=True)

        # 摘要标签页
        summary_frame = ttk.Frame(notebook)
        notebook.add(summary_frame, text="📊 摘要")

        self.summary_text = scrolledtext.ScrolledText(summary_frame, height=12, wrap=tk.WORD,
                                                     font=("Consolas", 9))
        self.summary_text.pack(fill=tk.BOTH, expand=True)

        # 详细结果标签页
        detail_frame = ttk.Frame(notebook)
        notebook.add(detail_frame, text="📄 详细结果")

        self.detail_text = scrolledtext.ScrolledText(detail_frame, height=12, wrap=tk.WORD,
                                                    font=("Consolas", 9))
        self.detail_text.pack(fill=tk.BOTH, expand=True)

        # 错误信息标签页
        error_frame = ttk.Frame(notebook)
        notebook.add(error_frame, text="⚠️ 错误信息")

        # 错误信息文本区域
        error_text_frame = ttk.Frame(error_frame)
        error_text_frame.pack(fill=tk.BOTH, expand=True)

        self.error_text = scrolledtext.ScrolledText(error_text_frame, height=12, wrap=tk.WORD,
                                                   font=("Consolas", 9))
        self.error_text.pack(fill=tk.BOTH, expand=True)

        # 编辑按钮区域
        self.edit_buttons_frame = ttk.Frame(error_frame)
        self.edit_buttons_frame.pack(fill=tk.X, pady=5)

        # 状态栏
        status_bar = ttk.Frame(main_frame)
        status_bar.pack(fill=tk.X, pady=(10, 0))

        self.status_bar_var = tk.StringVar(value="就绪 - 请选择输入目录并点击开始检查")
        status_bar_label = ttk.Label(status_bar, textvariable=self.status_bar_var,
                                    font=("Arial", 8), foreground="#888888")
        status_bar_label.pack(side=tk.LEFT)

        # 时间显示
        self.time_var = tk.StringVar()
        time_label = ttk.Label(status_bar, textvariable=self.time_var,
                              font=("Arial", 8), foreground="#888888")
        time_label.pack(side=tk.RIGHT)
        self.update_time()

    def update_time(self):
        """更新时间显示"""
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.time_var.set(current_time)
        self.root.after(1000, self.update_time)

    def load_last_directories(self):
        """加载上次使用的目录"""
        try:
            if os.path.exists("last_directories.json"):
                with open("last_directories.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.last_input_dir = data.get("input_dir", "")
                    self.last_output_dir = data.get("output_dir", "")
        except Exception as e:
            logger.warning(f"加载上次目录失败: {e}")

    def save_last_directories(self):
        """保存当前使用的目录"""
        try:
            data = {
                "input_dir": self.input_dir_var.get(),
                "output_dir": self.output_dir_var.get()
            }
            with open("last_directories.json", "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"保存目录配置失败: {e}")

    def browse_input_dir(self):
        """浏览输入目录"""
        initial_dir = self.last_input_dir if self.last_input_dir else os.path.expanduser("~")
        directory = filedialog.askdirectory(title="选择输入目录", initialdir=initial_dir)
        if directory:
            self.input_dir_var.set(directory)
            self.last_input_dir = directory
            self.save_last_directories()

    def browse_output_dir(self):
        """浏览输出目录"""
        initial_dir = self.last_output_dir if self.last_output_dir else os.path.expanduser("~")
        directory = filedialog.askdirectory(title="选择输出目录", initialdir=initial_dir)
        if directory:
            self.output_dir_var.set(directory)
            self.last_output_dir = directory
            self.save_last_directories()

    def update_progress(self, current, total, message, current_file="", phase="", estimated_time=""):
        """更新进度条 - 增强版"""
        if total > 0:
            progress = (current / total) * 100
            self.progress_var.set(progress)

        # 更新状态显示
        self.status_var.set(message)
        self.status_bar_var.set(f"正在处理: {message}")

        # 更新文件计数
        if total > 0:
            self.file_count_var.set(f"文件: {current}/{total}")

        # 更新详细信息
        if current_file:
            self.current_file_var.set(current_file)
        if phase:
            self.current_phase_var.set(phase)
        if estimated_time:
            self.estimated_time_var.set(estimated_time)

        # 更新内存使用情况
        try:
            import psutil
            process = psutil.Process()
            memory_mb = process.memory_info().rss / 1024 / 1024
            self.memory_usage_var.set(f"{memory_mb:.1f} MB")
        except ImportError:
            self.memory_usage_var.set("--")

        # 计算预估剩余时间
        if self.start_time and current > 0 and total > 0:
            elapsed_time = time.time() - self.start_time
            if current > 0:
                avg_time_per_file = elapsed_time / current
                remaining_files = total - current
                estimated_remaining = avg_time_per_file * remaining_files

                if estimated_remaining > 60:
                    time_str = f"{estimated_remaining/60:.1f} 分钟"
                elif estimated_remaining > 1:
                    time_str = f"{estimated_remaining:.0f} 秒"
                else:
                    time_str = "小于1秒"

                self.estimated_time_var.set(time_str)

        self.root.update_idletasks()

    def start_check(self):
        """开始检查"""
        input_dir = self.input_dir_var.get().strip()
        output_dir = self.output_dir_var.get().strip()

        if not input_dir:
            messagebox.showerror("错误", "请选择输入目录")
            return

        if not os.path.exists(input_dir):
            messagebox.showerror("错误", "输入目录不存在")
            return

        # 检查输入目录是否包含SHP文件
        shp_files = list(Path(input_dir).rglob("*.shp"))
        if not shp_files:
            if not messagebox.askyesno("警告", "输入目录中没有找到SHP文件，是否继续检查？"):
                return

        # 禁用按钮
        self.check_button.config(state=tk.DISABLED)
        self.export_button.config(state=tk.DISABLED)

        # 重置进度和状态
        self.progress_var.set(0)
        self.status_var.set("准备开始检查...")
        self.status_bar_var.set("正在初始化检查...")
        self.file_count_var.set("文件: 0/0")

        # 重置详细信息
        self.current_file_var.set("无")
        self.current_phase_var.set("无")
        self.estimated_time_var.set("--")
        self.memory_usage_var.set("--")

        # 记录开始时间
        self.start_time = time.time()
        self.end_time = None

        # 清空结果显示
        self.summary_text.delete(1.0, tk.END)
        self.detail_text.delete(1.0, tk.END)
        self.error_text.delete(1.0, tk.END)

        # 在新线程中运行检查
        def run_check():
            try:
                logger.info(f"开始检查目录: {input_dir}")

                # 获取几何修复配置
                auto_fix_geometry = self.auto_fix_geometry_var.get()
                geometry_tolerance = self.geometry_tolerance_var.get()

                logger.info(f"几何修复配置: 自动修复={auto_fix_geometry}, 容差={geometry_tolerance}")

                # 创建检查器实例，传入几何修复配置
                self.checker = GeoDataInspector(
                    input_dir,
                    output_dir,
                    self.field_config_manager,
                    auto_fix_geometry=auto_fix_geometry,
                    geometry_tolerance=geometry_tolerance
                )

                self.results = self.checker.run_check(progress_callback=self.update_progress)

                # 在主线程中更新UI
                self.root.after(0, self.update_results_display)
                logger.info("检查完成")

            except Exception as e:
                error_msg = str(e)
                logger.error(f"检查过程中出现错误: {error_msg}")

                # 使用用户友好的错误处理
                friendly_error = UserFriendlyErrorHandler.get_user_friendly_message(error_msg)
                self.root.after(0, lambda: messagebox.showerror("检查错误", friendly_error))
            finally:
                # 先设置结束时间，再调用check_completed
                self.root.after(0, lambda: setattr(self, 'end_time', time.time()))
                self.root.after(0, self.check_completed)

        thread = threading.Thread(target=run_check)
        thread.daemon = True
        thread.start()

    def check_completed(self):
        """检查完成后的处理"""
        # 记录结束时间
        self.end_time = time.time()

        self.check_button.config(state=tk.NORMAL)
        self.export_button.config(state=tk.NORMAL)
        self.progress_var.set(100)
        self.status_var.set("检查完成")

        # 计算并显示总用时
        if self.start_time and self.end_time:
            total_time = self.end_time - self.start_time
            if total_time > 60:
                time_str = f"{total_time/60:.1f} 分钟"
            elif total_time > 1:
                time_str = f"{total_time:.1f} 秒"
            else:
                time_str = f"{total_time*1000:.0f} 毫秒"

            self.status_bar_var.set(f"检查已完成，总用时: {time_str}，可以查看结果或导出报告")
        else:
            self.status_bar_var.set("检查已完成，可以查看结果或导出报告")

        if self.results:
            # 显示完成统计
            summary = self.results.get('summary', {})
            total_files = summary.get('total_files', 0)
            error_files = summary.get('error_files', 0)

            # 检查是否有字段合规性问题
            field_compliance_issues = []
            if self.results and isinstance(self.results, dict):
                for file_result in self.results.get('files', []) if isinstance(self.results.get('files', []), list) else []:
                    for field in file_result.get('fields', []):
                        compliance_issues = field.get('compliance_issues', [])
                        if isinstance(compliance_issues, list) and compliance_issues:
                            # 确保所有元素都是字符串
                            for issue in compliance_issues:
                                if isinstance(issue, str):
                                    field_compliance_issues.append(issue)

            # 计算检查用时
            check_time_str = ""
            if self.start_time:
                # 确保end_time已设置
                end_time = self.end_time if self.end_time else time.time()
                total_time = end_time - self.start_time
                if total_time > 60:
                    check_time_str = f"{total_time/60:.1f} 分钟"
                elif total_time > 1:
                    check_time_str = f"{total_time:.1f} 秒"
                else:
                    check_time_str = f"{total_time*1000:.0f} 毫秒"

            if error_files == 0 and not field_compliance_issues:
                messagebox.showinfo("完成", f"检查完成！\n\n共检查 {total_files} 个文件\n检查用时: {check_time_str}\n没有发现错误")
            else:
                issue_count = error_files + len(field_compliance_issues)
                messagebox.showwarning("完成", f"检查完成！\n\n共检查 {total_files} 个文件\n检查用时: {check_time_str}\n发现 {issue_count} 个问题\n请查看详细结果")
        else:
            messagebox.showwarning("完成", "检查完成，但没有生成结果")

    def update_results_display(self):
        """更新结果显示"""
        if not self.results:
            return

        # 更新摘要
        summary = self.results['summary']

        # 计算检查用时
        check_time_str = ""
        if self.start_time:
            # 如果end_time还没有设置，使用当前时间
            end_time = self.end_time if self.end_time else time.time()
            total_time = end_time - self.start_time
            if total_time > 60:
                check_time_str = f"{total_time/60:.1f} 分钟"
            elif total_time > 1:
                check_time_str = f"{total_time:.1f} 秒"
            else:
                check_time_str = f"{total_time*1000:.0f} 毫秒"

        summary_text = f"""检查结果摘要
{'='*50}
检查时间: {summary['check_time']}
检查用时: {check_time_str}
检查文件总数: {summary['total_files']}
SHP文件数量: {summary['shp_files']}
DBF文件数量: {summary['dbf_files']}
错误文件数量: {summary['error_files']}
总要素数量: {summary['total_features']}
总字段数量: {summary['total_fields']}

新增检查结果:
拓扑问题数量: {summary.get('topology_issues', 0)}
属性问题数量: {summary.get('attribute_issues', 0)}
基础问题数量: {summary.get('basic_issues', 0)}

几何修复信息:
几何修复记录: {len(self.results.get('geometry_fixes', []))} 个文件
"""
        self.summary_text.delete(1.0, tk.END)
        self.summary_text.insert(1.0, summary_text)

        # 更新详细结果
        detail_text = "详细检查结果\n" + "="*50 + "\n\n"
        files = self.results.get('files', []) if self.results and isinstance(self.results, dict) else []
        for file_result in files:
            detail_text += f"文件: {file_result['file_name']}\n"
            detail_text += f"路径: {file_result['file_path']}\n"

            if file_result.get('error'):
                detail_text += f"错误: {file_result['error']}\n"
            else:
                detail_text += f"几何类型: {file_result.get('geometry_type', 'N/A')}\n"
                detail_text += f"要素数量: {file_result.get('feature_count', 0)}\n"
                detail_text += f"字段数量: {file_result.get('field_count', 0)}\n"
                detail_text += f"文件大小: {file_result.get('file_size', 0)} bytes\n"

                # 字段信息
                if file_result.get('fields'):
                    detail_text += "字段信息:\n"
                    for field in file_result.get('fields', []):
                        detail_text += f"  {field['name']} ({field['type']})\n"
                        detail_text += f"    空值数量: {field['null_count']}\n"
                        detail_text += f"    唯一值数量: {field['unique_count']}\n"

                        if field.get('compliance_issues'):
                            detail_text += f"    合规性问题: {', '.join(field['compliance_issues'])}\n"

                # 新增检查结果
                if file_result.get('topology_issues'):
                    detail_text += "拓扑问题:\n"
                    for issue in file_result['topology_issues']:
                        detail_text += f"  {issue.get('type', '未知')}: {issue}\n"

                if file_result.get('attribute_issues'):
                    detail_text += "属性问题:\n"
                    for issue in file_result['attribute_issues']:
                        detail_text += f"  {issue.get('type', '未知')}: {issue.get('error', '未知错误')}\n"

                if file_result.get('basic_issues'):
                    detail_text += "基础问题:\n"
                    for issue in file_result['basic_issues']:
                        detail_text += f"  {issue.get('type', '未知')}: {issue.get('error', '未知错误')}\n"

                # 几何修复信息
                if file_result.get('geometry_fixes'):
                    detail_text += "几何修复信息:\n"
                    fix_info = file_result['geometry_fixes']
                    detail_text += f"  修复几何数量: {fix_info.get('fixed_count', 0)} 个\n"
                    detail_text += f"  备份文件: {fix_info.get('backup_path', 'N/A')}\n"
                    detail_text += f"  修复时间: {fix_info.get('fix_time', 'N/A')}\n"

            detail_text += "\n" + "-"*50 + "\n\n"

        self.detail_text.delete(1.0, tk.END)
        self.detail_text.insert(1.0, detail_text)

        # 优化错误信息显示
        self.update_error_display(files)

    def update_error_display(self, files):
        """优化错误信息显示 - 增强版"""
        error_text = "错误信息分类\n" + "="*50 + "\n\n"

        # 收集所有错误信息
        critical_errors = []  # 不可忽略错误
        ignorable_errors = []  # 可忽略错误
        field_edit_info = {}  # 按字段分组的编辑信息
        geometry_edit_info = {}  # 按文件分组的几何编辑信息

        # 错误统计
        error_stats = {
            'critical': 0,
            'high': 0,
            'medium': 0,
            'low': 0,
            'total': 0
        }

        # 处理字段合规性问题
        for file_result in files:
            file_name = file_result.get('file_name', '')
            file_path = file_result.get('file_path', '')

            # 处理文件级错误
            if file_result.get('error'):
                error_msg = file_result['error']
                error_type = UserFriendlyErrorHandler.classify_error(error_msg)
                error_priority = UserFriendlyErrorHandler.get_error_priority(error_type)

                error_info = {
                    'file_name': file_name,
                    'file_path': file_path,
                    'type': error_type,
                    'priority': error_priority,
                    'message': error_msg,
                    'friendly_message': UserFriendlyErrorHandler.get_user_friendly_message(error_msg, file_name)
                }

                if error_priority in ['critical', 'high']:
                    critical_errors.append(error_info)
                else:
                    ignorable_errors.append(error_info)

                error_stats[error_priority] += 1
                error_stats['total'] += 1

            # 处理几何问题
            if file_result.get('topology_issues'):
                geometry_key = f"{file_name}_geometry"
                geometry_edit_info[geometry_key] = {
                    'file_path': file_path,
                    'layer_name': None,  # 对于SHP文件，图层名为None
                    'issues': file_result['topology_issues'],
                    'level': 'medium'
                }

            # GDB多图层
            if file_result.get('layers'):
                for layer in file_result['layers']:
                    layer_name = layer.get('layer_name', '')

                    # 处理图层的几何问题
                    if layer.get('topology_issues'):
                        geometry_key = f"{file_name}_{layer_name}_geometry"
                        geometry_edit_info[geometry_key] = {
                            'file_path': file_path,
                            'layer_name': layer_name,
                            'issues': layer['topology_issues'],
                            'level': 'medium'
                        }

                    for field in layer.get('fields', []):
                        compliance_issues = field.get('compliance_issues', [])
                        if isinstance(compliance_issues, list) and compliance_issues:
                            field_name = field['name']
                            field_key = f"{file_name}_{layer_name}_{field_name}"

                            # 确定错误等级
                            error_level = get_field_error_level(field_name, file_name)

                            # 简化错误信息
                            simplified_issues = []
                            for issue in compliance_issues:
                                if "必填字段" in issue:
                                    simplified_issues.append("必填字段为空")
                                elif "字段类型" in issue:
                                    simplified_issues.append("字段类型不匹配")
                                elif "字段长度" in issue:
                                    simplified_issues.append("字段长度超限")
                                elif "数值范围" in issue:
                                    simplified_issues.append("数值范围异常")
                                elif "编码格式" in issue:
                                    simplified_issues.append("编码格式错误")
                                elif "数据完整性" in issue:
                                    simplified_issues.append("数据不完整")
                                elif "逻辑一致性" in issue:
                                    simplified_issues.append("逻辑不一致")
                                elif "空间参考" in issue:
                                    simplified_issues.append("空间参考不一致")
                                elif "字段值一致性" in issue:
                                    simplified_issues.append("字段值不一致")
                                else:
                                    simplified_issues.append(issue)

                            error_info = {
                                'file_name': file_name,
                                'layer_name': layer_name,
                                'field_name': field_name,
                                'issues': simplified_issues,
                                'null_count': field.get('null_count', 0),
                                'unique_count': field.get('unique_count', 0),
                                'level': error_level
                            }

                            if error_level == ERROR_LEVELS['CRITICAL']:
                                critical_errors.append(error_info)
                            else:
                                ignorable_errors.append(error_info)

                            # 存储编辑信息（按字段分组）
                            if field_key not in field_edit_info:
                                field_edit_info[field_key] = {
                                    'file_path': file_path,
                                    'field_name': field_name,
                                    'layer_name': layer_name,
                                    'issues': simplified_issues,
                                    'level': error_level
                                }

            # 普通SHP/DBF
            else:
                for field in file_result.get('fields', []):
                    compliance_issues = field.get('compliance_issues', [])
                    if isinstance(compliance_issues, list) and compliance_issues:
                        field_name = field['name']
                        field_key = f"{file_name}_{field_name}"

                        # 确定错误等级
                        error_level = get_field_error_level(field_name, file_name)

                        # 简化错误信息
                        simplified_issues = []
                        for issue in compliance_issues:
                            if "必填字段" in issue:
                                simplified_issues.append("必填字段为空")
                            elif "字段类型" in issue:
                                simplified_issues.append("字段类型不匹配")
                            elif "字段长度" in issue:
                                simplified_issues.append("字段长度超限")
                            elif "数值范围" in issue:
                                simplified_issues.append("数值范围异常")
                            elif "编码格式" in issue:
                                simplified_issues.append("编码格式错误")
                            elif "数据完整性" in issue:
                                simplified_issues.append("数据不完整")
                            elif "逻辑一致性" in issue:
                                simplified_issues.append("逻辑不一致")
                            elif "空间参考" in issue:
                                simplified_issues.append("空间参考不一致")
                            elif "字段值一致性" in issue:
                                simplified_issues.append("字段值不一致")
                            else:
                                simplified_issues.append(issue)

                        error_info = {
                            'file_name': file_name,
                            'field_name': field_name,
                            'issues': simplified_issues,
                            'null_count': field.get('null_count', 0),
                            'unique_count': field.get('unique_count', 0),
                            'level': error_level
                        }

                        if error_level == ERROR_LEVELS['CRITICAL']:
                            critical_errors.append(error_info)
                        else:
                            ignorable_errors.append(error_info)

                        # 存储编辑信息（按字段分组）
                        if field_key not in field_edit_info:
                            field_edit_info[field_key] = {
                                'file_path': file_path,
                                'field_name': field_name,
                                'layer_name': None,
                                'issues': simplified_issues,
                                'level': error_level
                            }

        # 显示不可忽略错误
        if critical_errors:
            error_text += "🚨 不可忽略错误 (必须修复):\n"
            error_text += "-" * 30 + "\n"
            for error in critical_errors:
                if 'file_name' in error:
                    error_text += f"📁 {error['file_name']}"
                    if error.get('layer_name'):
                        error_text += f" (图层: {error['layer_name']})"
                    error_text += f"\n  字段: {error.get('field_name', 'N/A')}\n"
                    error_text += f"  问题: {', '.join(error.get('issues', []))}\n"
                    if 'null_count' in error and 'unique_count' in error:
                        error_text += f"  空值: {error['null_count']}, 唯一值: {error['unique_count']}\n"
                    error_text += "\n"
                else:
                    # 处理文件级错误
                    error_text += f"📁 {error.get('file', 'N/A')}\n"
                    error_text += f"  错误类型: {error.get('type', 'N/A')}\n"
                    error_text += f"  错误信息: {error.get('message', 'N/A')}\n"
                    error_text += f"  解决建议: {error.get('friendly_message', 'N/A')}\n\n"

        # 显示可忽略错误
        if ignorable_errors:
            error_text += "⚠️ 可忽略错误 (建议修复):\n"
            error_text += "-" * 30 + "\n"
            for error in ignorable_errors:
                if 'file_name' in error:
                    error_text += f"📁 {error['file_name']}"
                    if error.get('layer_name'):
                        error_text += f" (图层: {error['layer_name']})"
                    error_text += f"\n  字段: {error.get('field_name', 'N/A')}\n"
                    error_text += f"  问题: {', '.join(error.get('issues', []))}\n"
                    if 'null_count' in error and 'unique_count' in error:
                        error_text += f"  空值: {error['null_count']}, 唯一值: {error['unique_count']}\n"
                    error_text += "\n"
                else:
                    # 处理文件级错误
                    error_text += f"📁 {error.get('file', 'N/A')}\n"
                    error_text += f"  错误类型: {error.get('type', 'N/A')}\n"
                    error_text += f"  错误信息: {error.get('message', 'N/A')}\n"
                    error_text += f"  解决建议: {error.get('friendly_message', 'N/A')}\n\n"

        # 处理其他错误类型
        errors = self.results.get('errors', []) if self.results and isinstance(self.results, dict) else []
        topology_issues = self.results.get('topology_issues', []) if self.results else []
        attribute_issues = self.results.get('attribute_issues', []) if self.results else []
        basic_issues = self.results.get('basic_issues', []) if self.results else []

        # 处理属性问题，添加到字段编辑信息中
        for issue in attribute_issues:
            if isinstance(issue, dict):
                file_path = issue.get('file', '')
                if file_path:
                    file_name = Path(file_path).name
                    # 为属性问题创建字段编辑信息
                    field_key = f"{file_name}_attribute"
                    field_edit_info[field_key] = {
                        'file_path': file_path,
                        'field_name': 'DLBM',  # 根据问题类型确定字段名
                        'layer_name': None,
                        'issues': [str(issue.get('issue', ''))],
                        'level': 'medium'
                    }

        if errors:
            error_text += "🚨 文件读取错误:\n"
            error_text += "-" * 30 + "\n"
            for error in errors:
                error_text += f"📁 {Path(error['file']).name}\n"
                error_text += f"  错误: {error['error']}\n\n"

        if topology_issues and isinstance(topology_issues, list):
            error_text += "⚠️ 拓扑问题:\n"
            error_text += "-" * 30 + "\n"
            for issue in topology_issues:
                if isinstance(issue, dict):
                    error_text += f"📁 {Path(str(issue.get('file', ''))).name}\n"
                    error_text += f"  问题: {str(issue.get('issue', ''))}\n\n"

        if attribute_issues and isinstance(attribute_issues, list):
            error_text += "⚠️ 属性问题:\n"
            error_text += "-" * 30 + "\n"
            for issue in attribute_issues:
                if isinstance(issue, dict):
                    file_name = str(issue.get('file', ''))
                    if file_name:
                        file_name = Path(file_name).name
                    error_text += f"📁 {file_name}\n"
                    issue_text = str(issue.get('issue', ''))
                    # 如果是字典格式的问题，提取错误信息
                    if isinstance(issue_text, dict):
                        error_type = issue_text.get('type', '')
                        error_msg = issue_text.get('error', '')
                        issue_text = f"{error_type}: {error_msg}"
                    error_text += f"  问题: {issue_text}\n\n"

        if basic_issues and isinstance(basic_issues, list):
            error_text += "⚠️ 基础问题:\n"
            error_text += "-" * 30 + "\n"
            for issue in basic_issues:
                if isinstance(issue, dict):
                    error_text += f"📁 {Path(str(issue.get('file', ''))).name}\n"
                    error_text += f"  问题: {str(issue.get('issue', ''))}\n\n"

        if not any([critical_errors, ignorable_errors, errors, topology_issues, attribute_issues, basic_issues]):
            error_text += "✅ 没有发现错误。\n"

        self.error_text.delete(1.0, tk.END)
        self.error_text.insert(1.0, error_text)

        # 创建优化的编辑按钮
        # 添加调试信息
        logger.info(f"字段编辑信息: {len(field_edit_info)} 个")
        logger.info(f"几何编辑信息: {len(geometry_edit_info)} 个")
        self.create_optimized_edit_buttons(field_edit_info, geometry_edit_info)

    def create_optimized_edit_buttons(self, field_edit_info, geometry_edit_info=None):
        """创建优化的编辑按钮 - 整合所有修复功能到同一级别"""
        # 清除现有的编辑按钮框架
        for widget in self.edit_buttons_frame.winfo_children():
            widget.destroy()

        if not field_edit_info and not geometry_edit_info:
            return

        # 分类字段编辑信息
        critical_fields = {}
        normal_fields = {}

        if field_edit_info:
            for key, info in field_edit_info.items():
                level = info.get('level', 'medium')
                if level == ERROR_LEVELS['CRITICAL']:
                    critical_fields[key] = info
                else:
                    normal_fields[key] = info

        # 创建必要修复按钮框架
        if critical_fields:
            critical_frame = ttk.LabelFrame(self.edit_buttons_frame, text="🚨 必要修复", padding="5")
            critical_frame.pack(fill=tk.X, padx=5, pady=2)

            # 必要修复的字段编辑按钮
            if FieldEditorDialog and critical_fields:
                critical_field_button = ttk.Button(
                    critical_frame,
                    text=f"字段编辑 ({len(critical_fields)}个)",
                    command=lambda: self.open_field_editor_dialog(critical_fields, "必要修复")
                )
                critical_field_button.pack(side=tk.LEFT, padx=5, pady=2)

        # 创建建议修复按钮框架
        if normal_fields or geometry_edit_info:
            normal_frame = ttk.LabelFrame(self.edit_buttons_frame, text="⚠️ 建议修复", padding="5")
            normal_frame.pack(fill=tk.X, padx=5, pady=2)

            # 建议修复的字段编辑按钮
            if FieldEditorDialog and normal_fields:
                normal_field_button = ttk.Button(
                    normal_frame,
                    text=f"字段编辑 ({len(normal_fields)}个)",
                    command=lambda: self.open_field_editor_dialog(normal_fields, "建议修复")
                )
                normal_field_button.pack(side=tk.LEFT, padx=5, pady=2)

            # 几何编辑按钮
            if GeometryEditorDialog and geometry_edit_info:
                geometry_button = ttk.Button(
                    normal_frame,
                    text=f"几何编辑 ({len(geometry_edit_info)}个)",
                    command=lambda: self.open_geometry_editor_dialog(geometry_edit_info)
                )
                geometry_button.pack(side=tk.LEFT, padx=5, pady=2)

        # 创建批量修复按钮框架 - 新增功能
        batch_frame = ttk.LabelFrame(self.edit_buttons_frame, text="🔧 批量修复", padding="5")
        batch_frame.pack(fill=tk.X, padx=5, pady=2)

        # 几何修复按钮
        geometry_fix_button = ttk.Button(
            batch_frame,
            text="🔧 几何修复",
            command=self.fix_geometry_only
        )
        geometry_fix_button.pack(side=tk.LEFT, padx=5, pady=2)

        # 缝隙修复按钮
        gap_fix_button = ttk.Button(
            batch_frame,
            text="🔗 缝隙修复",
            command=self.fix_gaps_only
        )
        gap_fix_button.pack(side=tk.LEFT, padx=5, pady=2)

        # 综合修复按钮
        comprehensive_fix_button = ttk.Button(
            batch_frame,
            text="🛠️ 综合修复",
            command=self.comprehensive_fix
        )
        comprehensive_fix_button.pack(side=tk.LEFT, padx=5, pady=2)

        # 添加说明标签
        info_text = "🚨 必要修复: 必须修复的错误\n⚠️ 建议修复: 可忽略但建议修复的错误\n🔧 批量修复: 对整个目录进行批量修复"
        info_label = ttk.Label(
            self.edit_buttons_frame,
            text=info_text,
            font=("Arial", 9),
            foreground="#666666"
        )
        info_label.pack(pady=5)

    def open_field_editor_dialog(self, edit_info, category="字段编辑"):
        """打开字段编辑器选择对话框"""
        try:
            if FieldEditorDialog is None:
                messagebox.showerror("错误", "字段编辑功能未启用，请确保field_editor_dialog.py文件存在")
                return

            if not isinstance(edit_info, dict) or not edit_info:
                messagebox.showerror("错误", "没有可编辑的字段")
                return

            # 创建字段选择对话框
            dialog = tk.Toplevel(self.root)
            dialog.title(f"{category} - 选择字段")
            dialog.geometry("800x600")  # 增加窗口大小
            dialog.minsize(700, 500)    # 设置最小窗口大小
            dialog.transient(self.root)
            dialog.grab_set()

            # 设置对话框位置
            dialog.update_idletasks()
            x = (dialog.winfo_screenwidth() // 2) - (800 // 2)
            y = (dialog.winfo_screenheight() // 2) - (600 // 2)
            dialog.geometry(f"800x600+{x}+{y}")

            # 确保对话框显示在最前面
            dialog.lift()
            dialog.focus_force()

            logger.info(f"字段编辑器对话框已创建，位置: ({x}, {y})")

            # 创建主框架
            main_frame = ttk.Frame(dialog, padding="10")
            main_frame.pack(fill=tk.BOTH, expand=True)

            # 标题
            title_label = ttk.Label(main_frame, text=f"{category} - 请选择要编辑的字段",
                                   font=("Arial", 12, "bold"))
            title_label.pack(pady=(0, 10))

            # 创建字段列表
            list_frame = ttk.Frame(main_frame)
            list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

            # 创建Treeview来显示字段信息
            columns = ('文件', '图层', '字段', '问题', '等级')
            tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=12)  # 减少表格高度，为按钮留出空间

            # 设置列标题
            for col in columns:
                tree.heading(col, text=col)

            # 设置列宽
            tree.column('文件', width=150)
            tree.column('图层', width=100)
            tree.column('字段', width=100)
            tree.column('问题', width=200)
            tree.column('等级', width=80)

            # 添加滚动条
            scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=tree.yview)
            tree.configure(yscrollcommand=scrollbar.set)

            # 布局
            tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

            # 填充字段信息
            field_items = []
            for key, info in edit_info.items():
                file_name = Path(info.get('file_path', '')).name
                layer_name = info.get('layer_name', 'N/A')
                field_name = info.get('field_name', 'N/A')
                issues = info.get('issues', [])
                level = info.get('level', 'medium')

                # 格式化问题信息
                issue_text = '; '.join(str(i) for i in issues) if issues else '无'

                # 确定等级显示
                level_display = "🚨 严重" if level == ERROR_LEVELS['CRITICAL'] else "⚠️ 一般"

                item = tree.insert('', 'end', values=(file_name, layer_name, field_name, issue_text, level_display))
                field_items.append((key, info))

            # 按钮框架
            button_frame = ttk.Frame(main_frame)
            button_frame.pack(fill=tk.X, pady=(10, 0))

            def open_selected_field():
                """打开选中的字段编辑器"""
                selection = tree.selection()
                if not selection:
                    messagebox.showwarning("警告", "请选择一个字段")
                    return

                # 获取选中的字段信息
                selected_item = tree.item(selection[0])
                selected_index = tree.index(selection[0])
                key, info = field_items[selected_index]

                file_path = info.get('file_path')
                field_name = info.get('field_name')
                layer_name = info.get('layer_name')

                # 添加调试信息
                logger.info(f"选中字段信息: file_path={file_path}, field_name={field_name}, layer_name={layer_name}")
                logger.info(f"字段信息详情: {info}")

                if not file_path or not field_name:
                    messagebox.showerror("错误", "无法获取文件路径或字段名")
                    logger.error(f"字段信息不完整: file_path={file_path}, field_name={field_name}")
                    return

                # 关闭选择对话框
                dialog.destroy()

                # 打开字段编辑弹窗
                if FieldEditorDialog is not None:
                    try:
                        logger.info(f"准备创建字段编辑器: file_path={file_path}, field_name={field_name}, layer_name={layer_name}")

                        # 验证参数
                        if not file_path or not os.path.exists(file_path):
                            raise FileNotFoundError(f"文件不存在: {file_path}")
                        if not field_name:
                            raise ValueError("字段名不能为空")

                        editor = FieldEditorDialog(
                            self.root,
                            file_path,
                            field_name,
                            layer_name
                        )

                        logger.info("字段编辑器创建成功")

                        # 等待编辑完成
                        if editor is not None and editor.run():
                            # 如果文件被修改，提示重新检查
                            if messagebox.askyesno("提示", "文件已修改，是否重新检查？"):
                                self.start_check()
                    except FileNotFoundError as e:
                        messagebox.showerror("错误", f"文件不存在: {str(e)}")
                        logger.error(f"文件不存在: {e}")
                    except ValueError as e:
                        messagebox.showerror("错误", f"参数错误: {str(e)}")
                        logger.error(f"参数错误: {e}")
                    except Exception as e:
                        messagebox.showerror("错误", f"打开字段编辑器失败: {str(e)}")
                        logger.error(f"打开字段编辑器失败: {e}", exc_info=True)
                else:
                    messagebox.showerror("错误", "字段编辑功能未启用")

            def open_all_fields():
                """批量打开所有字段编辑器"""
                if messagebox.askyesno("确认", f"确定要依次打开所有 {len(field_items)} 个字段编辑器吗？"):
                    dialog.destroy()

                    for i, (key, info) in enumerate(field_items):
                        file_path = info.get('file_path')
                        field_name = info.get('field_name')
                        layer_name = info.get('layer_name')

                        if not file_path or not field_name:
                            continue

                        # 显示进度
                        self.status_var.set(f"正在编辑字段 {i+1}/{len(field_items)}: {field_name}")
                        self.root.update()

                        # 打开字段编辑弹窗
                        editor = None
                        if FieldEditorDialog is not None:
                            try:
                                logger.info(f"批量编辑 - 准备创建字段编辑器: file_path={file_path}, field_name={field_name}, layer_name={layer_name}")

                                # 验证参数
                                if not file_path or not os.path.exists(file_path):
                                    raise FileNotFoundError(f"文件不存在: {file_path}")
                                if not field_name:
                                    raise ValueError("字段名不能为空")

                                editor = FieldEditorDialog(
                                    self.root,
                                    file_path,
                                    field_name,
                                    layer_name
                                )

                                logger.info("批量编辑 - 字段编辑器创建成功")

                                # 等待编辑完成
                                if editor is not None and editor.run():  # type: ignore
                                    # 如果文件被修改，询问是否继续
                                    if not messagebox.askyesno("提示", f"字段 {field_name} 已修改，是否继续编辑下一个字段？"):
                                        break
                            except FileNotFoundError as e:
                                messagebox.showerror("错误", f"文件不存在: {str(e)}")
                                logger.error(f"批量编辑 - 文件不存在: {e}")
                                continue
                            except ValueError as e:
                                messagebox.showerror("错误", f"参数错误: {str(e)}")
                                logger.error(f"批量编辑 - 参数错误: {e}")
                                continue
                            except Exception as e:
                                messagebox.showerror("错误", f"打开字段编辑器失败: {str(e)}")
                                logger.error(f"批量编辑 - 打开字段编辑器失败: {e}", exc_info=True)
                                continue
                        else:
                            messagebox.showerror("错误", "字段编辑功能未启用")
                            break

                    # 编辑完成后提示重新检查
                    if messagebox.askyesno("提示", "字段编辑已完成，是否重新检查？"):
                        self.start_check()

            # 按钮
            logger.info(f"创建字段编辑按钮，字段数量: {len(field_items)}")

            # 左侧按钮
            left_button_frame = ttk.Frame(button_frame)
            left_button_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

            # 右侧按钮
            right_button_frame = ttk.Frame(button_frame)
            right_button_frame.pack(side=tk.RIGHT, fill=tk.X)

            # 创建按钮
            edit_button = ttk.Button(left_button_frame, text="编辑选中字段", command=open_selected_field)
            edit_button.pack(side=tk.LEFT, padx=(0, 10))

            batch_button = ttk.Button(left_button_frame, text="批量编辑所有字段", command=open_all_fields)
            batch_button.pack(side=tk.LEFT, padx=(0, 10))

            cancel_button = ttk.Button(right_button_frame, text="取消", command=dialog.destroy)
            cancel_button.pack(side=tk.RIGHT)

            # 添加按钮状态调试
            logger.info(f"按钮创建完成: 编辑按钮={edit_button}, 批量按钮={batch_button}, 取消按钮={cancel_button}")

            # 添加调试信息
            logger.info(f"字段编辑器对话框创建完成，包含 {len(field_items)} 个字段")
            for i, (key, info) in enumerate(field_items):
                logger.info(f"字段 {i+1}: {info.get('field_name', 'N/A')} - {info.get('file_path', 'N/A')}")

        except Exception as e:
            messagebox.showerror("错误", f"打开字段编辑器失败: {str(e)}")
            logger.error(f"打开字段编辑器失败: {e}")

    def open_field_editor(self, edit_info):
        """打开字段编辑器（兼容旧版本）"""
        self.open_field_editor_dialog(edit_info, "字段编辑")

    def export_report(self):
        """导出报告"""
        if not self.results:
            messagebox.showwarning("警告", "没有可导出的结果")
            return

        # 选择导出格式
        format_var = tk.StringVar(value="word")
        format_dialog = tk.Toplevel(self.root)
        format_dialog.title("选择导出格式")
        format_dialog.geometry("350x200")
        format_dialog.transient(self.root)
        format_dialog.grab_set()

        ttk.Label(format_dialog, text="请选择导出格式:").pack(pady=10)

        ttk.Radiobutton(format_dialog, text="Word格式 (.docx) - 正式报告", variable=format_var, value="word").pack()
        ttk.Radiobutton(format_dialog, text="Excel格式 (.xlsx) - 详细数据", variable=format_var, value="excel").pack()
        ttk.Radiobutton(format_dialog, text="JSON格式 (.json) - 原始数据", variable=format_var, value="json").pack()

        def do_export():
            try:
                if self.checker is None:
                    messagebox.showerror("错误", "没有可导出的检查结果")
                    return
                output_path = self.checker.save_results(format_var.get())
                messagebox.showinfo("成功", f"报告已导出到:\n{output_path}")
                format_dialog.destroy()
            except Exception as e:
                messagebox.showerror("错误", f"导出失败: {str(e)}")

        ttk.Button(format_dialog, text="导出", command=do_export).pack(pady=10)
        ttk.Button(format_dialog, text="取消", command=format_dialog.destroy).pack()

    def open_field_config(self):
        """打开字段配置对话框 - 使用PandasTable"""
        try:
            from pandastable_field_config import FieldConfigPandasTable

            # 获取当前的重要文件配置
            current_critical_files = self.get_critical_files_config()

            dialog = FieldConfigPandasTable(
                self.root,
                default_standards=DEFAULT_FIELD_STANDARDS,
                critical_files_config=current_critical_files
            )
            dialog.run()

            # 获取完整配置并更新到主程序
            complete_config = dialog.get_complete_config()
            self.field_config_manager.update_field_standards(complete_config["field_standards"])

            # 更新重要文件配置
            self.update_critical_files_config(complete_config["critical_files"])

        except Exception as e:
            messagebox.showerror("错误", f"打开字段配置失败: {str(e)}")
            messagebox.showinfo("提示", "请确保已安装pandastable: pip install pandastable")

    def get_critical_files_config(self):
        """获取当前重要文件配置"""
        # 默认配置
        return {
            "GHMC": ["YDFW", "GHJX"],  # 规划名称字段在YDFW或GHJX文件中为不可忽略
            "PFDATE": ["YDFW", "GHJX"]  # 批准日期字段在YDFW或GHJX文件中为不可忽略
        }

    def update_critical_files_config(self, new_config):
        """更新重要文件配置"""
        try:
            # 这里可以将配置保存到文件或内存中
            # 目前先保存到实例变量中
            self.critical_files_config = new_config
            logger.info(f"已更新重要文件配置: {new_config}")
        except Exception as e:
            logger.error(f"更新重要文件配置失败: {e}")

    def get_field_error_level(self, field_name, file_name):
        """根据字段名和文件名确定错误等级 - 使用用户配置"""
        try:
            # 获取当前重要文件配置
            critical_files_config = getattr(self, 'critical_files_config', self.get_critical_files_config())

            # 转换为大写以便比较
            file_name_upper = file_name.upper()

            # 检查是否为特殊字段
            if field_name in critical_files_config:
                required_patterns = critical_files_config[field_name]
                for pattern in required_patterns:
                    if pattern.upper() in file_name_upper:
                        return ERROR_LEVELS['CRITICAL']  # 不可忽略

            # 默认返回可忽略
            return ERROR_LEVELS['IGNORABLE']

        except Exception as e:
            logger.error(f"获取字段错误等级失败: {e}")
            return ERROR_LEVELS['IGNORABLE']

    def clear_results(self):
        """清空结果"""
        if messagebox.askyesno("确认清空", "确定要清空所有检查结果吗？"):
            self.results = None
            self.checker = None
            self.summary_text.delete(1.0, tk.END)
            self.detail_text.delete(1.0, tk.END)
            self.error_text.delete(1.0, tk.END)
            self.progress_var.set(0)
            self.status_var.set("就绪")
            self.status_bar_var.set("就绪 - 请选择输入目录并点击开始检查")
            self.file_count_var.set("文件: 0")
            self.export_button.config(state=tk.DISABLED)
            logger.info("已清空检查结果")

    def fix_geometry_only(self):
        """仅修复几何错误，不进行其他检查"""
        input_dir = self.input_dir_var.get().strip()

        if not input_dir:
            messagebox.showerror("错误", "请选择输入目录")
            return

        if not os.path.exists(input_dir):
            messagebox.showerror("错误", "输入目录不存在")
            return

        # 确认操作
        if not messagebox.askyesno("确认修复",
                                  "确定要修复目录中所有文件的几何错误吗？\n\n"
                                  "注意：修复前会自动创建备份文件(.backup)"):
            return

        # 禁用按钮
        self.check_button.config(state=tk.DISABLED)

        # 重置进度和状态
        self.progress_var.set(0)
        self.status_var.set("准备开始修复几何...")
        self.status_bar_var.set("正在初始化几何修复...")

        # 记录开始时间
        self.start_time = time.time()
        self.end_time = None

        # 在新线程中运行几何修复
        def run_geometry_fix():
            try:
                logger.info(f"开始修复目录几何错误: {input_dir}")

                # 获取几何修复配置
                geometry_tolerance = self.geometry_tolerance_var.get()

                # 创建检查器实例
                self.checker = GeoDataInspector(
                    input_dir,
                    self.output_dir_var.get().strip() or str(Path.cwd()),
                    self.field_config_manager,
                    auto_fix_geometry=False,  # 这里设为False，因为我们要手动调用修复
                    geometry_tolerance=geometry_tolerance
                )

                # 执行几何修复
                fix_results = self.checker.auto_fix_all_geometry_files(
                    progress_callback=self.update_progress
                )

                # 在主线程中更新UI
                self.root.after(0, lambda: self.show_geometry_fix_results(fix_results))
                logger.info("几何修复完成")

            except Exception as e:
                error_msg = str(e)
                logger.error(f"几何修复过程中出现错误: {error_msg}")

                # 使用用户友好的错误处理
                friendly_error = UserFriendlyErrorHandler.get_user_friendly_message(error_msg)
                self.root.after(0, lambda: messagebox.showerror("修复错误", friendly_error))
            finally:
                # 恢复按钮状态
                self.root.after(0, lambda: setattr(self, 'end_time', time.time()))
                self.root.after(0, self.geometry_fix_completed)

        thread = threading.Thread(target=run_geometry_fix)
        thread.daemon = True
        thread.start()

    def fix_gaps_only(self):
        """仅修复缝隙错误，不进行其他检查"""
        input_dir = self.input_dir_var.get().strip()

        if not input_dir:
            messagebox.showerror("错误", "请选择输入目录")
            return

        if not os.path.exists(input_dir):
            messagebox.showerror("错误", "输入目录不存在")
            return

        # 确认操作
        if not messagebox.askyesno("确认修复",
                                  "确定要修复目录中所有文件的缝隙错误吗？\n\n"
                                  "注意：修复前会自动创建备份文件(.backup)"):
            return

        # 禁用按钮
        self.check_button.config(state=tk.DISABLED)

        # 重置进度和状态
        self.progress_var.set(0)
        self.status_var.set("准备开始修复缝隙...")
        self.status_bar_var.set("正在初始化缝隙修复...")

        # 记录开始时间
        self.start_time = time.time()
        self.end_time = None

        # 在新线程中运行缝隙修复
        def run_gap_fix():
            try:
                logger.info(f"开始修复目录缝隙错误: {input_dir}")

                # 获取几何修复配置
                geometry_tolerance = self.geometry_tolerance_var.get()

                # 创建检查器实例
                self.checker = GeoDataInspector(
                    input_dir,
                    self.output_dir_var.get().strip() or str(Path.cwd()),
                    self.field_config_manager,
                    auto_fix_geometry=False,
                    geometry_tolerance=geometry_tolerance
                )

                # 执行缝隙修复
                fix_results = self.checker.auto_fix_all_gaps_files(
                    progress_callback=self.update_progress
                )

                # 在主线程中更新UI
                self.root.after(0, lambda: self.show_gap_fix_results(fix_results))
                logger.info("缝隙修复完成")

            except Exception as e:
                error_msg = str(e)
                logger.error(f"缝隙修复过程中出现错误: {error_msg}")

                # 使用用户友好的错误处理
                friendly_error = UserFriendlyErrorHandler.get_user_friendly_message(error_msg)
                self.root.after(0, lambda: messagebox.showerror("修复错误", friendly_error))
            finally:
                # 恢复按钮状态
                self.root.after(0, lambda: setattr(self, 'end_time', time.time()))
                self.root.after(0, self.gap_fix_completed)

        thread = threading.Thread(target=run_gap_fix)
        thread.daemon = True
        thread.start()

    def comprehensive_fix(self):
        """综合修复：同时修复几何错误和缝隙错误"""
        input_dir = self.input_dir_var.get().strip()

        if not input_dir:
            messagebox.showerror("错误", "请选择输入目录")
            return

        if not os.path.exists(input_dir):
            messagebox.showerror("错误", "输入目录不存在")
            return

        # 确认操作
        if not messagebox.askyesno("确认修复",
                                  "确定要进行综合修复吗？\n\n"
                                  "将依次修复：\n"
                                  "1. 几何错误\n"
                                  "2. 缝隙错误\n\n"
                                  "注意：修复前会自动创建备份文件(.backup)"):
            return

        # 禁用按钮
        self.check_button.config(state=tk.DISABLED)

        # 重置进度和状态
        self.progress_var.set(0)
        self.status_var.set("准备开始综合修复...")
        self.status_bar_var.set("正在初始化综合修复...")

        # 记录开始时间
        self.start_time = time.time()
        self.end_time = None

        # 在新线程中运行综合修复
        def run_comprehensive_fix():
            try:
                logger.info(f"开始综合修复目录: {input_dir}")

                # 获取几何修复配置
                geometry_tolerance = self.geometry_tolerance_var.get()

                # 创建检查器实例
                self.checker = GeoDataInspector(
                    input_dir,
                    self.output_dir_var.get().strip() or str(Path.cwd()),
                    self.field_config_manager,
                    auto_fix_geometry=False,
                    geometry_tolerance=geometry_tolerance
                )

                # 执行综合修复
                fix_results = self.checker.comprehensive_fix_all_files(
                    progress_callback=self.update_progress
                )

                # 在主线程中更新UI
                self.root.after(0, lambda: self.show_comprehensive_fix_results(fix_results))
                logger.info("综合修复完成")

            except Exception as e:
                error_msg = str(e)
                logger.error(f"综合修复过程中出现错误: {error_msg}")

                # 使用用户友好的错误处理
                friendly_error = UserFriendlyErrorHandler.get_user_friendly_message(error_msg)
                self.root.after(0, lambda: messagebox.showerror("修复错误", friendly_error))
            finally:
                # 恢复按钮状态
                self.root.after(0, lambda: setattr(self, 'end_time', time.time()))
                self.root.after(0, self.comprehensive_fix_completed)

        thread = threading.Thread(target=run_comprehensive_fix)
        thread.daemon = True
        thread.start()

    def show_geometry_fix_results(self, fix_results):
        """显示几何修复结果"""
        if not fix_results or not fix_results.get('success'):
            error_msg = fix_results.get('error', '未知错误') if fix_results else '未知错误'
            messagebox.showerror("修复失败", f"几何修复失败: {error_msg}")
            return

        # 显示修复结果摘要
        summary = f"""几何修复完成！

修复统计:
• 总文件数: {fix_results['total_files']}
• 成功修复: {fix_results['successful_files']} 个文件
• 修复失败: {fix_results['failed_files']} 个文件
• 总修复几何: {fix_results['total_fixed']} 个
• 总错误数: {fix_results['total_errors']} 个

详细信息已记录到日志中。"""

        messagebox.showinfo("修复完成", summary)

        # 更新状态
        self.status_var.set("几何修复完成")
        self.status_bar_var.set(f"几何修复完成，共修复 {fix_results['total_fixed']} 个几何错误")

    def geometry_fix_completed(self):
        """几何修复完成后的处理"""
        # 记录结束时间
        self.end_time = time.time()

        # 恢复按钮状态
        self.check_button.config(state=tk.NORMAL)
        self.progress_var.set(100)

        # 计算并显示总用时
        if self.start_time and self.end_time:
            total_time = self.end_time - self.start_time
            if total_time > 60:
                time_str = f"{total_time/60:.1f} 分钟"
            elif total_time > 1:
                time_str = f"{total_time:.1f} 秒"
            else:
                time_str = f"{total_time*1000:.0f} 毫秒"

            self.status_bar_var.set(f"几何修复已完成，总用时: {time_str}")
        else:
            self.status_bar_var.set("几何修复已完成")

    def show_gap_fix_results(self, fix_results):
        """显示缝隙修复结果"""
        if not fix_results or not fix_results.get('success'):
            error_msg = fix_results.get('error', '未知错误') if fix_results else '未知错误'
            messagebox.showerror("缝隙修复失败", f"缝隙修复过程中出现错误：\n{error_msg}")
            return

        # 显示修复结果
        total_files = fix_results.get('total_files', 0)
        processed_files = fix_results.get('processed_files', 0)
        total_gaps = fix_results.get('total_gaps', 0)
        repaired_gaps = fix_results.get('repaired_gaps', 0)

        result_message = f"缝隙修复完成！\n\n"
        result_message += f"处理文件: {processed_files}/{total_files}\n"
        result_message += f"发现缝隙: {total_gaps} 个\n"
        result_message += f"修复缝隙: {repaired_gaps} 个\n"

        if fix_results.get('backup_created'):
            result_message += f"\n已创建备份文件"

        messagebox.showinfo("缝隙修复结果", result_message)

    def gap_fix_completed(self):
        """缝隙修复完成后的回调"""
        # 恢复按钮状态
        self.check_button.config(state=tk.NORMAL)

        # 更新状态栏
        if self.end_time:
            duration = self.end_time - self.start_time
            self.status_bar_var.set(f"缝隙修复完成 - 耗时: {duration:.1f}秒")
        else:
            self.status_bar_var.set("缝隙修复完成")

    def show_comprehensive_fix_results(self, fix_results):
        """显示综合修复结果"""
        if not fix_results or not fix_results.get('success'):
            error_msg = fix_results.get('error', '未知错误') if fix_results else '未知错误'
            messagebox.showerror("综合修复失败", f"综合修复过程中出现错误：\n{error_msg}")
            return

        # 显示修复结果
        total_files = fix_results.get('total_files', 0)
        processed_files = fix_results.get('processed_files', 0)

        geometry_results = fix_results.get('geometry_results', {})
        gap_results = fix_results.get('gap_results', {})

        total_geometry_errors = geometry_results.get('total_errors', 0)
        fixed_geometry_errors = geometry_results.get('fixed_errors', 0)
        total_gaps = gap_results.get('total_gaps', 0)
        repaired_gaps = gap_results.get('repaired_gaps', 0)

        result_message = f"综合修复完成！\n\n"
        result_message += f"处理文件: {processed_files}/{total_files}\n\n"
        result_message += f"几何修复:\n"
        result_message += f"  发现错误: {total_geometry_errors} 个\n"
        result_message += f"  修复错误: {fixed_geometry_errors} 个\n\n"
        result_message += f"缝隙修复:\n"
        result_message += f"  发现缝隙: {total_gaps} 个\n"
        result_message += f"  修复缝隙: {repaired_gaps} 个\n"

        if fix_results.get('backup_created'):
            result_message += f"\n已创建备份文件"

        messagebox.showinfo("综合修复结果", result_message)

    def comprehensive_fix_completed(self):
        """综合修复完成后的回调"""
        # 恢复按钮状态
        self.check_button.config(state=tk.NORMAL)

        # 更新状态栏
        if self.end_time:
            duration = self.end_time - self.start_time
            self.status_bar_var.set(f"综合修复完成 - 耗时: {duration:.1f}秒")
        else:
            self.status_bar_var.set("综合修复完成")

    def run(self):
        """运行GUI"""
        self.root.mainloop()

    def open_geometry_editor_dialog(self, edit_info):
        """打开几何编辑器选择对话框"""
        if GeometryEditorDialog is None:
            messagebox.showerror("错误", "几何编辑模块未加载")
            return

        try:
            if not isinstance(edit_info, dict) or not edit_info:
                messagebox.showerror("错误", "没有可编辑的几何文件")
                return

            # 创建几何文件选择对话框
            dialog = tk.Toplevel(self.root)
            dialog.title("几何编辑 - 选择文件")
            dialog.geometry("800x600")
            dialog.transient(self.root)
            dialog.grab_set()

            # 设置对话框位置
            dialog.update_idletasks()
            x = (dialog.winfo_screenwidth() // 2) - (600 // 2)
            y = (dialog.winfo_screenheight() // 2) - (400 // 2)
            dialog.geometry(f"800x600+{x}+{y}")

            # 创建主框架
            main_frame = ttk.Frame(dialog, padding="10")
            main_frame.pack(fill=tk.BOTH, expand=True)

            # 标题
            title_label = ttk.Label(main_frame, text="几何编辑 - 请选择要编辑的文件",
                                   font=("Arial", 12, "bold"))
            title_label.pack(pady=(0, 10))

            # 创建文件列表
            list_frame = ttk.Frame(main_frame)
            list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

            # 创建Treeview来显示文件信息
            columns = ('文件', '图层', '问题', '等级')
            tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=15)

            # 设置列标题
            for col in columns:
                tree.heading(col, text=col)

            # 设置列宽
            tree.column('文件', width=200)
            tree.column('图层', width=150)
            tree.column('问题', width=200)
            tree.column('等级', width=80)

            # 添加滚动条
            scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=tree.yview)
            tree.configure(yscrollcommand=scrollbar.set)

            # 布局
            tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

            # 填充文件信息
            file_items = []
            for key, info in edit_info.items():
                file_name = Path(info.get('file_path', '')).name
                layer_name = info.get('layer_name', 'N/A')
                issues = info.get('issues', [])
                level = info.get('level', 'medium')

                # 格式化问题信息
                issue_text = '; '.join(str(i) for i in issues) if issues else '无'

                # 确定等级显示
                level_display = "🚨 严重" if level == ERROR_LEVELS['CRITICAL'] else "⚠️ 一般"

                item = tree.insert('', 'end', values=(file_name, layer_name, issue_text, level_display))
                file_items.append((key, info))

            # 按钮框架
            button_frame = ttk.Frame(main_frame)
            button_frame.pack(fill=tk.X, pady=(10, 0))

            def open_selected_file():
                """打开选中的几何编辑器"""
                selection = tree.selection()
                if not selection:
                    messagebox.showwarning("警告", "请选择一个文件")
                    return

                # 获取选中的文件信息
                selected_item = tree.item(selection[0])
                selected_index = tree.index(selection[0])
                key, info = file_items[selected_index]

                file_path = info.get('file_path')
                layer_name = info.get('layer_name')

                if not file_path:
                    messagebox.showerror("错误", "无法获取文件路径")
                    return

                # 关闭选择对话框
                dialog.destroy()

                # 打开几何编辑对话框
                if GeometryEditorDialog is not None:
                    geometry_dialog = GeometryEditorDialog(self.root, file_path, layer_name)
                else:
                    geometry_dialog = None
                has_changes = geometry_dialog is not None and geometry_dialog.run()  # type: ignore

                if has_changes:
                    messagebox.showinfo("完成", "几何编辑已完成")
                    # 可以在这里添加刷新结果的逻辑
                    if messagebox.askyesno("提示", "几何已修改，是否重新检查？"):
                        self.start_check()

            def open_all_files():
                """批量打开所有几何编辑器"""
                if messagebox.askyesno("确认", f"确定要依次打开所有 {len(file_items)} 个几何编辑器吗？"):
                    dialog.destroy()

                    for i, (key, info) in enumerate(file_items):
                        file_path = info.get('file_path')
                        layer_name = info.get('layer_name')

                        if not file_path:
                            continue

                        # 显示进度
                        self.status_var.set(f"正在编辑几何 {i+1}/{len(file_items)}: {Path(file_path).name}")
                        self.root.update()

                        # 打开几何编辑对话框
                        if GeometryEditorDialog is not None:
                            geometry_dialog = GeometryEditorDialog(self.root, file_path, layer_name)
                        else:
                            geometry_dialog = None
                        has_changes = geometry_dialog is not None and geometry_dialog.run()  # type: ignore

                        if has_changes:
                            # 如果文件被修改，询问是否继续
                            if not messagebox.askyesno("提示", f"几何 {Path(file_path).name} 已修改，是否继续编辑下一个文件？"):
                                break

                    # 编辑完成后提示重新检查
                    if messagebox.askyesno("提示", "几何编辑已完成，是否重新检查？"):
                        self.start_check()

            # 按钮
            ttk.Button(button_frame, text="编辑选中文件", command=open_selected_file).pack(side=tk.LEFT, padx=(0, 10))
            ttk.Button(button_frame, text="批量编辑所有文件", command=open_all_files).pack(side=tk.LEFT, padx=(0, 10))
            ttk.Button(button_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT)

        except Exception as e:
            logger.error(f"打开几何编辑器失败: {e}")
            messagebox.showerror("错误", f"打开几何编辑器失败: {str(e)}")

    def open_geometry_editor(self, edit_info):
        """打开几何编辑对话框（兼容旧版本）"""
        self.open_geometry_editor_dialog(edit_info)

def main():
    """主函数"""
    app = GeoDataInspectorGUI()
    app.run()

if __name__ == "__main__":
    main()
